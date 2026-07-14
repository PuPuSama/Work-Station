from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from config import AppConfig
from models import TaskRecord, TdkMetadata
from services.article_images import extract_article_title
from services.article_validation import strip_llm_code_fence
from services.generator import primary_keyword, render_prompt
from services.llm import LLMClient
from storage import now_iso


DESCRIPTION_MAX_CHARACTERS = 150
KEYWORD_COUNT = 6
TDK_PROMPT_VERSION = "tdk-v1"
WORD_FONT_NAME = "Times New Roman"


class TdkGenerationError(RuntimeError):
    """Raised when the model cannot produce metadata matching the hard contract."""


def current_article(task: TaskRecord) -> str:
    return (
        task.final_article
        or task.linked_article
        or task.humanized_article
        or task.initial_article
        or task.article
    ).strip()


def article_title(task: TaskRecord, article: str) -> str:
    return (extract_article_title(article) or task.selected_title or task.topic).strip()


def _json_object(text: str) -> dict[str, Any]:
    cleaned = strip_llm_code_fence(text).strip()
    try:
        payload = json.loads(cleaned)
    except json.JSONDecodeError:
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start < 0 or end <= start:
            raise TdkGenerationError("The model did not return valid TDK JSON.") from None
        try:
            payload = json.loads(cleaned[start : end + 1])
        except json.JSONDecodeError as exc:
            raise TdkGenerationError("The model did not return valid TDK JSON.") from exc
    if not isinstance(payload, dict):
        raise TdkGenerationError("The model TDK response must be a JSON object.")
    return payload


def parse_tdk_response(text: str, *, title: str, article: str) -> TdkMetadata:
    payload = _json_object(text)
    description = re.sub(r"\s+", " ", str(payload.get("description") or "")).strip()
    raw_keywords = payload.get("keywords")
    if isinstance(raw_keywords, str):
        keywords = [item.strip() for item in raw_keywords.split(",")]
    elif isinstance(raw_keywords, list):
        keywords = [str(item).strip() for item in raw_keywords]
    else:
        keywords = []
    keywords = [keyword for keyword in keywords if keyword]

    if not description:
        raise TdkGenerationError("The SEO description is empty.")
    if len(description) > DESCRIPTION_MAX_CHARACTERS:
        raise TdkGenerationError(
            f"The SEO description has {len(description)} characters; maximum is "
            f"{DESCRIPTION_MAX_CHARACTERS}."
        )
    if not re.search(r"[A-Za-z]", description):
        raise TdkGenerationError("The SEO description must be in English.")
    if len(keywords) != KEYWORD_COUNT:
        raise TdkGenerationError(
            f"The SEO metadata must contain exactly {KEYWORD_COUNT} keywords."
        )
    if len({keyword.casefold() for keyword in keywords}) != KEYWORD_COUNT:
        raise TdkGenerationError("The SEO keywords must be distinct.")
    if any("," in keyword or "\n" in keyword for keyword in keywords):
        raise TdkGenerationError("Each SEO keyword must be a single comma-free phrase.")

    return TdkMetadata(
        title=title,
        description=description,
        keywords=keywords,
        description_character_count=len(description),
        source_article_hash=hashlib.sha256(article.encode("utf-8")).hexdigest(),
        generated_at=now_iso(),
        prompt_version=TDK_PROMPT_VERSION,
    )


def generate_tdk_metadata(
    config: AppConfig,
    task: TaskRecord,
    *,
    llm: LLMClient | None = None,
) -> TdkMetadata:
    article = current_article(task)
    if not article:
        raise TdkGenerationError("The current article is empty; TDK cannot be generated.")
    title = article_title(task, article)
    if not title:
        raise TdkGenerationError("The current article has no title.")

    client = llm or LLMClient(config)
    prompt = render_prompt(
        "tdk",
        TITLE=title,
        PRIMARY_KEYWORD=primary_keyword(task),
        ARTICLE=article,
    )
    last_error = ""
    for attempt in range(2):
        retry_note = (
            "\n\nYour previous response failed validation: "
            f"{last_error} Return corrected JSON only."
            if attempt and last_error
            else ""
        )
        result = client.chat(
            [
                {
                    "role": "system",
                    "content": "You are a senior B2B Google SEO metadata editor.",
                },
                {"role": "user", "content": prompt + retry_note},
            ],
            temperature=0.3,
            max_tokens=500,
        )
        if not result:
            raise TdkGenerationError("The language model returned no TDK content.")
        try:
            return parse_tdk_response(result, title=title, article=article)
        except TdkGenerationError as exc:
            last_error = str(exc)

    raise TdkGenerationError(last_error or "Unable to generate valid SEO TDK metadata.")


def _set_font(run, font_name: str, size_pt: float, *, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run_properties = run._element.get_or_add_rPr()
    run_fonts = run_properties.get_or_add_rFonts()
    run_fonts.set(qn("w:ascii"), font_name)
    run_fonts.set(qn("w:hAnsi"), font_name)
    run_fonts.set(qn("w:eastAsia"), font_name)


def export_tdk_docx(task: TaskRecord, metadata: TdkMetadata) -> Path:
    """Write the workflow's compact-reference TDK document as ``D.docx``."""

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = WORD_FONT_NAME
    normal.font.size = Pt(11)
    normal_rpr = normal.element.get_or_add_rPr()
    normal_fonts = normal_rpr.get_or_add_rFonts()
    normal_fonts.set(qn("w:ascii"), WORD_FONT_NAME)
    normal_fonts.set(qn("w:hAnsi"), WORD_FONT_NAME)
    normal_fonts.set(qn("w:eastAsia"), WORD_FONT_NAME)
    normal_fonts.set(qn("w:cs"), WORD_FONT_NAME)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    entries = (
        ("T", metadata.title),
        ("D", metadata.description),
        ("K", ", ".join(metadata.keywords)),
    )
    for label, value in entries:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.25
        _set_font(paragraph.add_run(f"{label}: "), WORD_FONT_NAME, 11, bold=True)
        _set_font(paragraph.add_run(value), WORD_FONT_NAME, 11)

    output_dir = Path(task.task_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "D.docx"
    document.save(output_path)
    return output_path
