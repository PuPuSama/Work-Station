from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from config import AppConfig


TEXT_EXTENSIONS = {".txt", ".md", ".csv"}
DOCX_EXTENSIONS = {".docx"}
XLSX_EXTENSIONS = {".xlsx"}
PDF_EXTENSIONS = {".pdf"}


def collect_customer_context(config: AppConfig, customer: str, max_chars: int = 14000) -> str:
    folder = config.knowledge_base / customer
    if not folder.exists():
        return ""

    chunks: list[str] = []
    for path in sorted(folder.rglob("*")):
        if not path.is_file() or path.name.startswith("~$"):
            continue
        text = extract_text(path)
        if text.strip():
            chunks.append(f"## Source: {path.name}\n{text.strip()}")
        joined = "\n\n".join(chunks)
        if len(joined) >= max_chars:
            return joined[:max_chars]
    return "\n\n".join(chunks)[:max_chars]


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix in TEXT_EXTENSIONS:
            return read_text_like(path)
        if suffix in DOCX_EXTENSIONS:
            return read_docx(path)
        if suffix in XLSX_EXTENSIONS:
            return read_xlsx(path)
        if suffix in PDF_EXTENSIONS:
            return read_pdf(path)
    except Exception as exc:
        return f"[Could not read {path.name}: {exc}]"
    return ""


def read_text_like(path: Path) -> str:
    if path.suffix.lower() == ".csv":
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            reader = csv.reader(handle)
            return "\n".join(" | ".join(row) for row in reader)
    return path.read_text(encoding="utf-8", errors="ignore")


def read_docx(path: Path) -> str:
    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_lines: list[str] = []
    for table in document.tables:
        for row in table.rows:
            table_lines.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(paragraphs + table_lines)


def read_xlsx(path: Path) -> str:
    workbook = load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        sheet.reset_dimensions()
        parts.append(f"Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value is not None]
            if values:
                parts.append(" | ".join(values))
    workbook.close()
    return "\n".join(parts)


def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages[:20]:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)
