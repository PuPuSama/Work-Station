from __future__ import annotations

import hashlib
import re
from collections import defaultdict
from collections.abc import Iterable, Mapping
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shape import CT_Inline
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.opc.part import Part
from docx.shared import Inches, Pt, RGBColor, Twips

from config import AppConfig
from models import TaskRecord
from services.article_images import (
    ImagePlacement,
    ImageValidationError,
    build_image_audit_markdown,
    image_pixel_size,
    prepare_task_images,
    resolve_asset_image_placements,
    resolve_image_placements,
    sanitize_image_stem,
)
from services.article_validation import validate_article_layout
from services.generator import enforce_homepage_brand_link


MARKDOWN_LINK_PATTERN = re.compile(r"\[([^\]]+)\]\((https?://[^)\s]+)\)")
URL_PATTERN = re.compile(r"(https?://[^\s)]+)")
BOLD_TEXT_PATTERN = re.compile(r"\*\*(.+?)\*\*")
FAQ_BOLD_QUESTION_PATTERN = re.compile(r"^\s*\*\*(Q:\s+.+)\*\*\s*$")
TABLE_SEPARATOR_CELL_PATTERN = re.compile(r"^:?-{3,}:?$")
WORD_FONT_NAME = "Times New Roman"
BLACK_HEX = "000000"
TABLE_HEADER_FILL_HEX = "E7E6E6"
TABLE_CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
TABLE_INDENT_DXA = TABLE_CELL_MARGINS_DXA["start"]


@dataclass(frozen=True, slots=True)
class EmbeddedArticleImage:
    """Verified private WebP bytes used by the Server DOCX renderer."""

    asset_id: str
    data: bytes
    filename: str
    width: int
    height: int


def _value(item: object, *names: str, default: Any = "") -> Any:
    for name in names:
        if isinstance(item, Mapping) and name in item:
            value = item[name]
        else:
            value = getattr(item, name, None)
        if value is not None and value != "":
            return value
    return default


def safe_filename(value: str) -> str:
    return sanitize_image_stem(value, fallback="article", max_length=120)


def _current_article(task: object, title: str) -> str:
    article = str(
        _value(
            task,
            "final_article",
            "linked_article",
            "humanized_article",
            "initial_article",
            "article",
            default="",
        )
        or ""
    )
    return article or f"# {title}\n\n"


def _has_image_sources(task: object) -> bool:
    if str(_value(task, "hero_image", default="") or "").strip():
        return True
    if list(_value(task, "images", default=[]) or []):
        return True
    return any(
        str(_value(product, "image_path", "source_path", default="") or "").strip()
        for product in list(_value(task, "products", default=[]) or [])
    )


def _prepared_images(task: object, markdown: str) -> list[object]:
    existing = list(_value(task, "images", default=[]) or [])
    if existing:
        paths = [
            str(_value(image, "prepared_path", "webp_path", default="") or "").strip()
            for image in existing
        ]
        has_hero = any(
            str(_value(image, "role", "kind", "type", default="product")).casefold() == "hero"
            for image in existing
        )
        if has_hero and all(
            path and Path(path).is_file() and Path(path).suffix.casefold() == ".webp"
            for path in paths
        ):
            return existing

    if not _has_image_sources(task):
        return []
    return prepare_task_images(task, markdown, require_hero=True)


def _validate_embedded_images(
    images: list[object],
    embedded_images: Mapping[str, EmbeddedArticleImage],
) -> None:
    expected_ids = {
        str(_value(image, "prepared_asset_id", default="") or "").strip()
        for image in images
    }
    if "" in expected_ids or expected_ids != set(embedded_images):
        raise ImageValidationError(
            "Server DOCX image bytes do not match the Task image assets."
        )
    for image in images:
        asset_id = str(
            _value(image, "prepared_asset_id", default="") or ""
        ).strip()
        payload = embedded_images[asset_id]
        expected_hash = str(
            _value(image, "prepared_content_hash", default="") or ""
        ).casefold()
        expected_filename = str(
            _value(image, "filename", default="") or ""
        )
        expected_width = _value(image, "width", default=None)
        expected_height = _value(image, "height", default=None)
        if (
            payload.asset_id != asset_id
            or hashlib.sha256(payload.data).hexdigest() != expected_hash
            or payload.filename != expected_filename
            or payload.width != expected_width
            or payload.height != expected_height
        ):
            raise ImageValidationError(
                "Server DOCX image bytes do not match the Task image assets."
            )


def _build_task_document(
    config: AppConfig,
    task: TaskRecord,
    *,
    embedded_images: Mapping[str, EmbeddedArticleImage] | None = None,
) -> tuple[Document, str, list[object]]:
    title = task.selected_title or task.topic or "Article"
    markdown = _current_article(task, title)
    markdown = enforce_homepage_brand_link(markdown, task)
    validate_article_layout(markdown)
    if embedded_images is None:
        images = _prepared_images(task, markdown)
    else:
        images = list(task.images)
        _validate_embedded_images(images, embedded_images)

    document = Document()
    configure_styles(document, config)
    render_markdown(
        document,
        markdown,
        config,
        images=images,
        embedded_images=embedded_images,
    )
    enforce_document_typography(document)
    return document, markdown, images


def build_task_docx_bytes(
    config: AppConfig,
    task: TaskRecord,
    *,
    embedded_images: Mapping[str, EmbeddedArticleImage],
) -> bytes:
    """Render a Server article DOCX without creating local task files."""

    document, _markdown, _images = _build_task_document(
        config,
        task,
        embedded_images=embedded_images,
    )
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def export_task_docx(config: AppConfig, task: TaskRecord) -> Path:
    document, markdown, images = _build_task_document(config, task)
    title = task.selected_title or task.topic or "Article"

    output_dir = Path(task.task_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    (output_dir / "07_final_with_images.md").write_text(
        build_image_audit_markdown(markdown, images),
        encoding="utf-8",
    )
    output_path = output_dir / f"{safe_filename(title)}.docx"
    document.save(output_path)
    return output_path


def configure_styles(document: Document, config: AppConfig) -> None:
    styles = document.styles
    set_style_font(styles["Normal"], WORD_FONT_NAME, config.body_size, bold=False)
    set_style_font(styles["Title"], WORD_FONT_NAME, config.title_1_size, bold=True)
    set_style_font(styles["Heading 1"], WORD_FONT_NAME, config.title_1_size, bold=True)
    set_style_font(styles["Heading 2"], WORD_FONT_NAME, config.title_2_size, bold=True)
    set_style_font(styles["Heading 3"], WORD_FONT_NAME, config.title_3_size, bold=True)

    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def set_style_font(style, font_name: str, size_pt: float, bold: bool) -> None:
    font = style.font
    font.name = font_name
    font.size = Pt(size_pt)
    font.bold = bold
    font.color.rgb = RGBColor(0, 0, 0)
    run_properties = style.element.get_or_add_rPr()
    run_fonts = run_properties.get_or_add_rFonts()
    run_fonts.set(qn("w:ascii"), font_name)
    run_fonts.set(qn("w:hAnsi"), font_name)
    run_fonts.set(qn("w:eastAsia"), font_name)


def _set_ooxml_run_font(run_element, *, black: bool = False) -> None:
    run_properties = run_element.get_or_add_rPr()
    run_fonts = run_properties.get_or_add_rFonts()
    run_fonts.set(qn("w:ascii"), WORD_FONT_NAME)
    run_fonts.set(qn("w:hAnsi"), WORD_FONT_NAME)
    run_fonts.set(qn("w:eastAsia"), WORD_FONT_NAME)
    run_fonts.set(qn("w:cs"), WORD_FONT_NAME)
    if black:
        color = run_properties.find(qn("w:color"))
        if color is None:
            color = OxmlElement("w:color")
            run_properties.append(color)
        color.set(qn("w:val"), BLACK_HEX)
        color.attrib.pop(qn("w:themeColor"), None)


def enforce_document_typography(document: Document) -> None:
    """Stamp Times New Roman everywhere and black on non-link heading runs."""

    for run_element in document.element.body.xpath(".//w:r"):
        _set_ooxml_run_font(run_element)

    heading_styles = {"Title", "Heading 1", "Heading 2", "Heading 3"}
    for paragraph in document.paragraphs:
        if paragraph.style and paragraph.style.name in heading_styles:
            for run_element in paragraph._p.xpath(".//w:r"):
                if run_element.getparent().tag == qn("w:hyperlink"):
                    continue
                _set_ooxml_run_font(run_element, black=True)


def _split_markdown_table_row(line: str) -> list[str] | None:
    stripped = line.strip()
    if "|" not in stripped:
        return None

    cells: list[str] = []
    current: list[str] = []
    delimiter_count = 0
    trailing_delimiter = False
    index = 0
    while index < len(stripped):
        character = stripped[index]
        if character == "\\" and index + 1 < len(stripped) and stripped[index + 1] == "|":
            current.append("|")
            trailing_delimiter = False
            index += 2
            continue
        if character == "|":
            cells.append("".join(current).strip())
            current = []
            delimiter_count += 1
            trailing_delimiter = True
        else:
            current.append(character)
            trailing_delimiter = False
        index += 1
    cells.append("".join(current).strip())

    if delimiter_count == 0:
        return None
    if stripped.startswith("|"):
        cells = cells[1:]
    if trailing_delimiter:
        cells = cells[:-1]
    return cells or None


def _table_alignment(separator: str):
    marker = separator.strip()
    if marker.startswith(":") and marker.endswith(":"):
        return WD_ALIGN_PARAGRAPH.CENTER
    if marker.endswith(":"):
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


def _parse_markdown_table(
    lines: list[str],
    start_index: int,
) -> tuple[list[list[str]], list[Any], int] | None:
    if start_index + 1 >= len(lines):
        return None

    header = _split_markdown_table_row(lines[start_index])
    separator = _split_markdown_table_row(lines[start_index + 1])
    if not header or not separator or len(header) != len(separator):
        return None
    if not all(TABLE_SEPARATOR_CELL_PATTERN.fullmatch(cell.strip()) for cell in separator):
        return None

    rows = [header]
    next_index = start_index + 2
    while next_index < len(lines) and lines[next_index].strip():
        row = _split_markdown_table_row(lines[next_index])
        if row is None:
            break
        if len(row) != len(header):
            return None
        rows.append(row)
        next_index += 1

    alignments = [_table_alignment(cell) for cell in separator]
    return rows, alignments, next_index


def _inline_display_text(text: str) -> str:
    without_links = MARKDOWN_LINK_PATTERN.sub(lambda match: match.group(1), text)
    return BOLD_TEXT_PATTERN.sub(lambda match: match.group(1), without_links)


def _table_column_widths(rows: list[list[str]], total_width_dxa: int) -> list[int]:
    column_count = len(rows[0])
    weights = []
    for column_index in range(column_count):
        visible_lengths = [
            len(_inline_display_text(row[column_index]).strip()) for row in rows
        ]
        weights.append(float(max(8, min(max(visible_lengths, default=0), 32))))

    total_weight = sum(weights)
    widths = [int(round(total_width_dxa * weight / total_weight)) for weight in weights]
    widths[-1] += total_width_dxa - sum(widths)
    return widths


def _ensure_ooxml_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _set_ooxml_width(parent, tag: str, width_dxa: int) -> None:
    width = _ensure_ooxml_child(parent, tag)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(int(width_dxa)))


def _set_table_cell_margins(cell) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = _ensure_ooxml_child(cell_properties, "w:tcMar")
    for side, width_dxa in TABLE_CELL_MARGINS_DXA.items():
        margin = _ensure_ooxml_child(margins, f"w:{side}")
        margin.set(qn("w:type"), "dxa")
        margin.set(qn("w:w"), str(width_dxa))


def _apply_table_geometry(document: Document, table, rows: list[list[str]]) -> None:
    section = document.sections[0]
    total_width_dxa = int(
        section.page_width.twips
        - section.left_margin.twips
        - section.right_margin.twips
    )
    widths = _table_column_widths(rows, total_width_dxa)

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_properties = table._tbl.tblPr
    _set_ooxml_width(table_properties, "w:tblW", total_width_dxa)

    indent = _ensure_ooxml_child(table_properties, "w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    layout = _ensure_ooxml_child(table_properties, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    table_grid = table._tbl.tblGrid
    for child in list(table_grid):
        table_grid.remove(child)
    for width_dxa in widths:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width_dxa))
        table_grid.append(grid_column)

    for column_index, width_dxa in enumerate(widths):
        table.columns[column_index].width = Twips(width_dxa)

    for row in table.rows:
        row.height = None
        for column_index, cell in enumerate(row.cells):
            width_dxa = widths[column_index]
            cell.width = Twips(width_dxa)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_ooxml_width(cell._tc.get_or_add_tcPr(), "w:tcW", width_dxa)
            _set_table_cell_margins(cell)


def add_markdown_table(
    document: Document,
    rows: list[list[str]],
    alignments: list[Any],
    config: AppConfig,
) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"

    for row_index, row_values in enumerate(rows):
        for column_index, value in enumerate(row_values):
            cell = table.cell(row_index, column_index)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = alignments[column_index]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            add_text_with_links(
                paragraph,
                value,
                config,
                bold=row_index == 0,
            )
            if row_index == 0:
                shading = _ensure_ooxml_child(cell._tc.get_or_add_tcPr(), "w:shd")
                shading.set(qn("w:fill"), TABLE_HEADER_FILL_HEX)

    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeating_header = _ensure_ooxml_child(header_properties, "w:tblHeader")
    repeating_header.set(qn("w:val"), "true")
    _apply_table_geometry(document, table, rows)


def render_markdown(
    document: Document,
    markdown: str,
    config: AppConfig,
    *,
    images: Iterable[object] | None = None,
    embedded_images: Mapping[str, EmbeddedArticleImage] | None = None,
) -> None:
    image_list = list(images or [])
    placements = (
        (
            resolve_asset_image_placements(markdown, image_list)
            if embedded_images is not None
            else resolve_image_placements(markdown, image_list)
        )
        if image_list
        else []
    )
    before: dict[int, list[ImagePlacement]] = defaultdict(list)
    after: dict[int, list[ImagePlacement]] = defaultdict(list)
    for placement in placements:
        (before if placement.position == "before" else after)[placement.line_index].append(placement)

    generated_markers = {placement.image["marker"] for placement in placements}
    lines = markdown.splitlines()

    def is_prose_line(index: int) -> bool:
        line = lines[index].rstrip()
        if not line or line.strip() in generated_markers:
            return False
        if "|" in line:
            return False
        if re.match(
            r"^(?:#{1,6}\s+|[-+*]\s+|\d+[.)]\s+|>\s*|```|~~~|---$|\*\*\*$|___$|!\[|img\.)",
            line.strip(),
            re.IGNORECASE,
        ):
            return False
        return _parse_markdown_table(lines, index) is None

    line_index = 0
    while line_index < len(lines):
        parsed_table = _parse_markdown_table(lines, line_index)
        if parsed_table is not None:
            table_rows, alignments, next_index = parsed_table
            for consumed_index in range(line_index, next_index):
                for placement in before.get(consumed_index, []):
                    add_inline_image(
                        document,
                        placement.image,
                        config,
                        embedded_images=embedded_images,
                    )
            add_markdown_table(document, table_rows, alignments, config)
            for consumed_index in range(line_index, next_index):
                for placement in after.get(consumed_index, []):
                    add_inline_image(
                        document,
                        placement.image,
                        config,
                        embedded_images=embedded_images,
                    )
            line_index = next_index
            continue

        for placement in before.get(line_index, []):
            add_inline_image(
                document,
                placement.image,
                config,
                embedded_images=embedded_images,
            )

        raw_line = lines[line_index]
        line = raw_line.rstrip()
        consumed_until = line_index + 1
        if is_prose_line(line_index):
            prose_lines = [line.strip()]
            while consumed_until < len(lines) and is_prose_line(consumed_until):
                prose_lines.append(lines[consumed_until].strip())
                consumed_until += 1

            for consumed_index in range(line_index + 1, consumed_until):
                for placement in before.get(consumed_index, []):
                    add_inline_image(
                        document,
                        placement.image,
                        config,
                        embedded_images=embedded_images,
                    )

            prose = " ".join(prose_lines)
            paragraph = document.add_paragraph()
            faq_question = FAQ_BOLD_QUESTION_PATTERN.fullmatch(prose)
            add_text_with_links(
                paragraph,
                faq_question.group(1) if faq_question else prose,
                config,
                bold=bool(faq_question),
            )
        elif line and line.strip() not in generated_markers:
            if line.startswith("# "):
                add_markdown_heading(document, line[2:].strip(), 1, config)
            elif line.startswith("## "):
                add_markdown_heading(document, line[3:].strip(), 2, config)
            elif line.startswith("### "):
                add_markdown_heading(document, line[4:].strip(), 3, config)
            elif line.startswith("- "):
                paragraph = document.add_paragraph(style="List Bullet")
                add_text_with_links(paragraph, line[2:].strip(), config)
            else:
                paragraph = document.add_paragraph()
                faq_question = FAQ_BOLD_QUESTION_PATTERN.fullmatch(line)
                add_text_with_links(
                    paragraph,
                    faq_question.group(1) if faq_question else line,
                    config,
                    bold=bool(faq_question),
                )

        after_line = [
            placement
            for consumed_index in range(line_index, consumed_until)
            for placement in after.get(consumed_index, [])
        ]
        if after_line and line and document.paragraphs:
            document.paragraphs[-1].paragraph_format.keep_with_next = True
        for placement in after_line:
            add_inline_image(
                document,
                placement.image,
                config,
                embedded_images=embedded_images,
            )
        line_index = consumed_until


def add_markdown_heading(
    document: Document,
    text: str,
    level: int,
    config: AppConfig,
) -> None:
    paragraph = document.add_heading(level=level)
    size_by_level = {
        1: config.title_1_size,
        2: config.title_2_size,
        3: config.title_3_size,
    }
    add_text_with_links(
        paragraph,
        text,
        config,
        bold=True,
        font_size_pt=size_by_level[level],
    )


def add_inline_image(
    document: Document,
    image: Mapping[str, Any],
    config: AppConfig,
    *,
    embedded_images: Mapping[str, EmbeddedArticleImage] | None = None,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    if embedded_images is None:
        add_webp_picture(run, Path(str(image["prepared_path"])))
    else:
        asset_id = str(image["prepared_asset_id"])
        payload = embedded_images.get(asset_id)
        if payload is None:
            raise ImageValidationError(
                "Server DOCX image bytes do not match the Task image assets."
            )
        add_webp_picture_bytes(
            run,
            data=payload.data,
            filename=payload.filename,
            width_px=payload.width,
            height_px=payload.height,
        )

    description = str(image.get("product_name") or image.get("filename") or "Article image")
    for doc_properties in paragraph._p.xpath(".//wp:docPr"):
        doc_properties.set("descr", description)
        doc_properties.set("title", description)

    marker_paragraph = document.add_paragraph()
    marker_paragraph.paragraph_format.space_before = Pt(0)
    if image.get("role") == "hero":
        marker_paragraph.paragraph_format.keep_with_next = True
    add_run(marker_paragraph, str(image["marker"]), config)


def add_webp_picture(run, path: Path) -> None:
    """Embed a WebP as an inline OOXML image.

    python-docx 1.1.x cannot parse WebP headers, so ``Run.add_picture`` rejects
    otherwise valid files. Creating an ordinary OPC image part plus DrawingML
    inline shape keeps the actual WebP bytes and remains standards-compliant.
    """

    width_px, height_px = image_pixel_size(path)
    add_webp_picture_bytes(
        run,
        data=path.read_bytes(),
        filename=path.name,
        width_px=width_px,
        height_px=height_px,
    )


def add_webp_picture_bytes(
    run,
    *,
    data: bytes,
    filename: str,
    width_px: int,
    height_px: int,
) -> None:
    """Embed already verified WebP bytes as an inline OOXML image part."""

    max_width_inches = 5.8
    max_height_inches = 6.4
    scale = min(max_width_inches / width_px, max_height_inches / height_px)
    width = Inches(width_px * scale)
    height = Inches(height_px * scale)

    package = run.part.package
    partname = package.next_partname("/word/media/image%d.webp")
    image_part = Part(partname, "image/webp", bytes(data), package)
    relationship_id = run.part.relate_to(
        image_part,
        RELATIONSHIP_TYPE.IMAGE,
    )
    inline = CT_Inline.new_pic_inline(
        run.part.next_id,
        relationship_id,
        filename,
        width,
        height,
    )
    run._r.add_drawing(inline)


def add_text_with_links(
    paragraph,
    text: str,
    config: AppConfig,
    *,
    bold: bool = False,
    font_size_pt: float | None = None,
) -> None:
    """Render Markdown bold spans, Markdown links, and bare URLs into real runs."""

    cursor = 0
    for match in BOLD_TEXT_PATTERN.finditer(text):
        if match.start() > cursor:
            _add_text_with_links_without_bold_markup(
                paragraph,
                text[cursor : match.start()],
                config,
                bold=bold,
                font_size_pt=font_size_pt,
            )
        _add_text_with_links_without_bold_markup(
            paragraph,
            match.group(1),
            config,
            bold=True,
            font_size_pt=font_size_pt,
        )
        cursor = match.end()
    if cursor < len(text):
        _add_text_with_links_without_bold_markup(
            paragraph,
            text[cursor:],
            config,
            bold=bold,
            font_size_pt=font_size_pt,
        )


def _add_text_with_links_without_bold_markup(
    paragraph,
    text: str,
    config: AppConfig,
    *,
    bold: bool = False,
    font_size_pt: float | None = None,
) -> None:
    cursor = 0
    for match in MARKDOWN_LINK_PATTERN.finditer(text):
        if match.start() > cursor:
            add_plain_text_with_links(
                paragraph,
                text[cursor : match.start()],
                config,
                bold=bold,
                font_size_pt=font_size_pt,
            )
        add_hyperlink(
            paragraph,
            match.group(1),
            match.group(2),
            config,
            bold=True,
            font_size_pt=font_size_pt,
        )
        cursor = match.end()
    if cursor < len(text):
        add_plain_text_with_links(
            paragraph,
            text[cursor:],
            config,
            bold=bold,
            font_size_pt=font_size_pt,
        )


def add_plain_text_with_links(
    paragraph,
    text: str,
    config: AppConfig,
    *,
    bold: bool = False,
    font_size_pt: float | None = None,
) -> None:
    cursor = 0
    for match in URL_PATTERN.finditer(text):
        if match.start() > cursor:
            add_run(
                paragraph,
                text[cursor : match.start()],
                config,
                bold=bold,
                font_size_pt=font_size_pt,
            )
        url = match.group(1).rstrip(".,;:")
        trailing = match.group(1)[len(url) :]
        add_hyperlink(
            paragraph,
            url,
            url,
            config,
            bold=True,
            font_size_pt=font_size_pt,
        )
        if trailing:
            add_run(
                paragraph,
                trailing,
                config,
                bold=bold,
                font_size_pt=font_size_pt,
            )
        cursor = match.end()
    if cursor < len(text):
        add_run(
            paragraph,
            text[cursor:],
            config,
            bold=bold,
            font_size_pt=font_size_pt,
        )


def add_run(
    paragraph,
    text: str,
    config: AppConfig,
    *,
    bold: bool = False,
    font_size_pt: float | None = None,
) -> None:
    if text:
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = WORD_FONT_NAME
        resolved_size = font_size_pt if font_size_pt is not None else config.body_size
        run.font.size = Pt(resolved_size)
        run_properties = run._element.get_or_add_rPr()
        run_fonts = run_properties.get_or_add_rFonts()
        run_fonts.set(qn("w:ascii"), WORD_FONT_NAME)
        run_fonts.set(qn("w:hAnsi"), WORD_FONT_NAME)
        run_fonts.set(qn("w:eastAsia"), WORD_FONT_NAME)
        run_fonts.set(qn("w:cs"), WORD_FONT_NAME)


def add_hyperlink(
    paragraph,
    text: str,
    url: str,
    config: AppConfig | None = None,
    *,
    bold: bool = True,
    font_size_pt: float | None = None,
) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    if bold:
        bold_element = OxmlElement("w:b")
        bold_element.set(qn("w:val"), "1")
        run_properties.append(bold_element)
        bold_complex = OxmlElement("w:bCs")
        bold_complex.set(qn("w:val"), "1")
        run_properties.append(bold_complex)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    if config is not None:
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), WORD_FONT_NAME)
        fonts.set(qn("w:hAnsi"), WORD_FONT_NAME)
        fonts.set(qn("w:eastAsia"), WORD_FONT_NAME)
        fonts.set(qn("w:cs"), WORD_FONT_NAME)
        run_properties.append(fonts)
        resolved_size = font_size_pt if font_size_pt is not None else config.body_size
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), str(int(round(resolved_size * 2))))
        run_properties.append(size)
        size_complex = OxmlElement("w:szCs")
        size_complex.set(qn("w:val"), str(int(round(resolved_size * 2))))
        run_properties.append(size_complex)

    run.append(run_properties)
    text_node = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace():
        text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
