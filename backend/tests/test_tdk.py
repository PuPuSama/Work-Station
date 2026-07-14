from __future__ import annotations

import sys
import tempfile
import unittest
from dataclasses import replace
from pathlib import Path

from docx import Document


BACKEND_DIR = Path(__file__).resolve().parents[1]
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from config import load_config  # noqa: E402
from models import STATUS_DOCX_EXPORTED, TaskRecord  # noqa: E402
from services.tdk import (  # noqa: E402
    TdkGenerationError,
    export_tdk_docx,
    generate_tdk_metadata,
    parse_tdk_response,
)


ARTICLE = """# Exact Original Article Title

Buyers often struggle to compare supplier claims with real application requirements.

## What Should Buyers Compare?

Review specifications, quality controls, and order requirements before requesting a quote.
"""


VALID_RESPONSE = """{
  "description": "Struggling to compare suppliers? Learn which specifications, quality checks, and order details lead to a safer sourcing decision.",
  "keywords": [
    "supplier comparison",
    "B2B sourcing guide",
    "product specifications",
    "quality control checklist",
    "purchase requirements",
    "supplier selection"
  ]
}"""


class StubLlm:
    def __init__(self, responses: list[str]):
        self.responses = list(responses)
        self.calls = 0

    def chat(self, messages, temperature=0.7, max_tokens=1800):
        self.calls += 1
        return self.responses.pop(0)


def task_at(path: Path) -> TaskRecord:
    return TaskRecord(
        id="tdk-test",
        week_folder="week",
        customer="example.com",
        topic_index=1,
        topic="Supplier comparison",
        status=STATUS_DOCX_EXPORTED,
        task_dir=str(path),
        selected_title="A Different Selected Title",
        final_article=ARTICLE,
        created_at="2026-07-10T00:00:00",
        updated_at="2026-07-10T00:00:00",
    )


class TdkGenerationTests(unittest.TestCase):
    def test_title_is_taken_exactly_from_the_article_h1(self) -> None:
        metadata = parse_tdk_response(
            VALID_RESPONSE,
            title="Exact Original Article Title",
            article=ARTICLE,
        )
        self.assertEqual(metadata.title, "Exact Original Article Title")
        self.assertLessEqual(metadata.description_character_count, 150)
        self.assertEqual(len(metadata.keywords), 6)

    def test_description_over_150_characters_is_rejected(self) -> None:
        response = (
            '{"description":"'
            + ("x" * 151)
            + '","keywords":["one","two","three","four","five","six"]}'
        )
        with self.assertRaises(TdkGenerationError):
            parse_tdk_response(response, title="Title", article=ARTICLE)

    def test_invalid_first_response_is_retried(self) -> None:
        cfg = replace(load_config(), output_root=Path("D:/unused"))
        llm = StubLlm(
            [
                '{"description":"' + ("x" * 151) + '","keywords":[]}',
                VALID_RESPONSE,
            ]
        )
        metadata = generate_tdk_metadata(cfg, task_at(Path("D:/unused")), llm=llm)
        self.assertEqual(llm.calls, 2)
        self.assertEqual(metadata.title, "Exact Original Article Title")

    def test_export_writes_exact_t_d_k_lines_to_d_docx(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            task = task_at(Path(directory))
            metadata = parse_tdk_response(
                VALID_RESPONSE,
                title="Exact Original Article Title",
                article=ARTICLE,
            )
            output = export_tdk_docx(task, metadata)
            self.assertEqual(output.name, "D.docx")
            document = Document(output)
            self.assertEqual(
                [paragraph.text for paragraph in document.paragraphs],
                [
                    "T: Exact Original Article Title",
                    f"D: {metadata.description}",
                    f"K: {', '.join(metadata.keywords)}",
                ],
            )
            section = document.sections[0]
            self.assertAlmostEqual(section.left_margin.inches, 1.0, places=2)
            self.assertEqual(document.styles["Normal"].font.name, "Times New Roman")
            self.assertTrue(
                all(
                    run.font.name == "Times New Roman"
                    for paragraph in document.paragraphs
                    for run in paragraph.runs
                )
            )


if __name__ == "__main__":
    unittest.main()
