from __future__ import annotations

import os
import sys
import tempfile
import unittest
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

import sqlalchemy as sa
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient


BACKEND_DIR = Path(__file__).resolve().parents[1]
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from knowledge_agent.http import router  # noqa: E402
from knowledge_agent.runtime import create_knowledge_runtime  # noqa: E402
from knowledge_agent import (  # noqa: E402
    EMBEDDING_DIMENSIONS,
    EmbeddingBatch,
    KnowledgeProject,
    WordPressProbeResult,
)


class FakeEmbeddingProvider:
    model_id = "test-embedding-model"
    dimensions = EMBEDDING_DIMENSIONS

    def __init__(self) -> None:
        self.calls: list[tuple[str, ...]] = []

    def embed(self, texts: tuple[str, ...]) -> EmbeddingBatch:
        self.calls.append(tuple(texts))
        vectors = tuple(
            (float(index + 1),) + (0.0,) * (EMBEDDING_DIMENSIONS - 1)
            for index, _text in enumerate(texts)
        )
        return EmbeddingBatch(vectors=vectors, model=self.model_id)


def docx_bytes() -> bytes:
    document = Document()
    document.core_properties.title = "Private Fastener Specification"
    document.add_heading("Wood Screw", level=1)
    document.add_paragraph("Material: carbon steel.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


class DisabledKnowledgeHttpTests(unittest.TestCase):
    def test_routes_return_not_found_without_enabled_runtime(self) -> None:
        app = FastAPI()
        app.include_router(router)
        with TestClient(app) as client:
            response = client.get("/api/knowledge/example.com")

        self.assertEqual(response.status_code, 404)
        self.assertEqual(response.json()["detail"], "Knowledge Agent is disabled.")


class FakeWordPressSync:
    def __init__(self) -> None:
        self.calls: list[str] = []

    def probe(self, site_url: str) -> WordPressProbeResult:
        self.calls.append(site_url)
        return WordPressProbeResult(
            site_url="https://example.com",
            detected=True,
            rest_api_url="https://example.com/wp-json/",
            namespaces=("wp/v2",),
            route_count=7,
            reason="WordPress REST index exposes the wp/v2 namespace or routes.",
        )


class WordPressHttpContractTests(unittest.TestCase):
    def setUp(self) -> None:
        self.sync = FakeWordPressSync()
        app = FastAPI()
        app.state.knowledge_agent_runtime = SimpleNamespace(
            wordpress_sync=self.sync
        )
        app.include_router(router)
        self.client = TestClient(app)

    def tearDown(self) -> None:
        self.client.close()

    def test_probe_returns_detection_evidence_without_exposing_credentials(self) -> None:
        response = self.client.post(
            "/api/knowledge/www.example.com/wordpress/probe",
            json={"site_url": "https://example.com"},
        )

        self.assertEqual(response.status_code, 200, response.text)
        payload = response.json()
        self.assertEqual(payload["project_id"], "example.com")
        self.assertTrue(payload["detected"])
        self.assertEqual(payload["namespaces"], ["wp/v2"])
        self.assertEqual(self.sync.calls, ["https://example.com"])
        self.assertNotIn("api_key", response.text.casefold())

    def test_probe_rejects_cross_project_site_before_network_call(self) -> None:
        response = self.client.post(
            "/api/knowledge/example.com/wordpress/probe",
            json={"site_url": "https://other.test"},
        )

        self.assertEqual(response.status_code, 422)
        self.assertEqual(self.sync.calls, [])


@unittest.skipUnless(
    os.environ.get("ARTICLE_AGENT_DATABASE_URL"),
    "ARTICLE_AGENT_DATABASE_URL is required for PostgreSQL integration tests",
)
class KnowledgeHttpIntegrationTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.temp_directory = tempfile.TemporaryDirectory()
        cls.embedding_provider = FakeEmbeddingProvider()
        cls.runtime = create_knowledge_runtime(
            database_url=os.environ["ARTICLE_AGENT_DATABASE_URL"],
            artifact_root=Path(cls.temp_directory.name),
            embedding_provider=cls.embedding_provider,
        )
        app = FastAPI()
        app.state.knowledge_agent_runtime = cls.runtime

        def enqueue_research(*, action, graph_request, approved_urls):
            if action != "start":
                raise ValueError("test queue supports start only")
            run = cls.runtime.research_execution.enqueue(graph_request)
            return {
                "run": run,
                "batch_id": "batch-research-1",
                "job_id": "job-research-1",
            }

        app.state.knowledge_research_enqueue = enqueue_research
        app.include_router(router)
        cls.client = TestClient(app)

    @classmethod
    def tearDownClass(cls) -> None:
        cls.client.close()
        cls.runtime.close()
        cls.temp_directory.cleanup()

    def setUp(self) -> None:
        with self.runtime.engine.begin() as connection:
            for table in (
                "gap_fill_attempts",
                "research_graph_events",
                "research_graph_runs",
                "evidence_links",
                "evidence_pack_hits",
                "evidence_packs",
                "retrieval_scopes",
                "retrieval_plans",
                "knowledge_product_asset_evidence",
                "knowledge_product_source_evidence",
                "knowledge_products",
                "snapshot_assets",
                "knowledge_assets",
                "knowledge_chunks",
                "source_snapshots",
                "knowledge_sources",
                "projects",
            ):
                connection.execute(
                    sa.text(f"DELETE FROM {table} WHERE project_id = :project_id"),
                    {"project_id": "example.com"},
                )

    def test_m4_research_run_is_queued_and_queryable(self) -> None:
        self.runtime.repository.upsert_project(
            KnowledgeProject(
                project_id="example.com",
                customer_name="Example",
                official_domain="example.com",
            )
        )
        plan = self.client.post(
            "/api/knowledge/example.com/retrieval-plans",
            json={
                "retrieval_plan_id": "plan-research-v1",
                "article_id": "topic_006",
                "outline_version": 1,
                "scopes": [
                    {
                        "scope_id": "scope-overview",
                        "ordinal": 0,
                        "scope_type": "h2_section",
                        "scope_key": "overview",
                        "title": "Overview",
                        "query_variants": ["fastener overview"],
                        "minimum_hits": 1,
                        "minimum_distinct_sources": 1,
                    }
                ],
            },
        )
        self.assertEqual(plan.status_code, 201, plan.text)

        created = self.client.post(
            "/api/knowledge/example.com/research-runs",
            json={
                "organization_id": "org-example",
                "retrieval_plan_id": "plan-research-v1",
                "max_discovery_queries": 2,
            },
        )
        self.assertEqual(created.status_code, 202, created.text)
        payload = created.json()
        self.assertEqual(payload["run"]["status"], "queued")
        self.assertEqual(payload["queue_batch_id"], "batch-research-1")
        thread_id = payload["run"]["thread_id"]

        listing = self.client.get(
            "/api/knowledge/example.com/research-runs",
            params={"article_id": "topic_006"},
        )
        self.assertEqual(listing.status_code, 200, listing.text)
        self.assertEqual([item["thread_id"] for item in listing.json()], [thread_id])

        detail = self.client.get(
            f"/api/knowledge/example.com/research-runs/{thread_id}"
        )
        self.assertEqual(detail.status_code, 200, detail.text)
        self.assertEqual(
            [event["event_type"] for event in detail.json()["events"]],
            ["queued"],
        )
        self.assertEqual(detail.json()["review_candidates"], [])

    def tearDown(self) -> None:
        self.setUp()

    def test_upload_list_and_open_original_evidence(self) -> None:
        content = docx_bytes()
        response = self.client.post(
            "/api/knowledge/www.example.com/sources/upload",
            files={
                "file": (
                    "private-spec.docx",
                    content,
                    (
                        "application/vnd.openxmlformats-officedocument."
                        "wordprocessingml.document"
                    ),
                )
            },
            data={
                "display_name": "Private specification",
                "trust_tier": "hard_fact",
            },
        )

        self.assertEqual(response.status_code, 201, response.text)
        uploaded = response.json()
        self.assertEqual(uploaded["project_id"], "example.com")
        self.assertEqual(uploaded["status"], "inbox")
        self.assertEqual(uploaded["parser_name"], "docx-lightweight")
        self.assertGreater(uploaded["chunk_count"], 0)

        repeated = self.client.post(
            "/api/knowledge/example.com/sources/upload",
            files={"file": ("private-spec.docx", content)},
            data={
                "display_name": "Private specification",
                "trust_tier": "hard_fact",
            },
        )
        self.assertEqual(repeated.status_code, 201, repeated.text)
        self.assertEqual(
            repeated.json()["snapshot_id"],
            uploaded["snapshot_id"],
        )

        listing = self.client.get("/api/knowledge/example.com")
        self.assertEqual(listing.status_code, 200, listing.text)
        library = listing.json()
        self.assertEqual(library["source_count"], 1)
        self.assertEqual(library["inbox_count"], 1)
        self.assertEqual(library["published_count"], 0)
        self.assertEqual(len(library["sources"]), 1)
        source = library["sources"][0]
        self.assertEqual(
            source["classification_reason"],
            "operator uploaded private document",
        )
        self.assertTrue(source["raw_evidence_url"])

        evidence = self.client.get(source["raw_evidence_url"])
        self.assertEqual(evidence.status_code, 200, evidence.text)
        self.assertEqual(evidence.content, content)

    def test_upload_rejects_unsupported_format_without_creating_source(self) -> None:
        response = self.client.post(
            "/api/knowledge/example.com/sources/upload",
            files={"file": ("notes.txt", b"plain text")},
        )

        self.assertEqual(response.status_code, 422)
        listing = self.client.get("/api/knowledge/example.com")
        self.assertEqual(listing.status_code, 200)
        self.assertEqual(listing.json()["source_count"], 0)

    def test_source_requires_review_before_embedding_and_atomic_publication(self) -> None:
        uploaded = self.client.post(
            "/api/knowledge/example.com/sources/upload",
            files={"file": ("private-spec.docx", docx_bytes())},
        ).json()
        source_id = uploaded["source_id"]

        premature = self.client.post(
            f"/api/knowledge/example.com/sources/{source_id}/publish"
        )
        self.assertEqual(premature.status_code, 409)
        self.assertIn("approved", premature.json()["detail"])

        reviewed = self.client.put(
            f"/api/knowledge/example.com/sources/{source_id}/review",
            json={
                "source_kind": "private_file",
                "trust_tier": "hard_fact",
                "decision": "approve",
                "reason": "Operator verified the supplied product specification.",
            },
        )
        self.assertEqual(reviewed.status_code, 200, reviewed.text)
        self.assertEqual(reviewed.json()["decision"], "approve")

        published = self.client.post(
            f"/api/knowledge/example.com/sources/{source_id}/publish"
        )
        self.assertEqual(published.status_code, 200, published.text)
        self.assertEqual(published.json()["status"], "published")
        self.assertEqual(
            published.json()["embedding_model"],
            self.embedding_provider.model_id,
        )
        self.assertGreater(published.json()["chunk_count"], 0)

        library = self.client.get("/api/knowledge/example.com").json()
        self.assertEqual(library["published_count"], 1)
        self.assertEqual(library["inbox_count"], 0)
        self.assertEqual(
            library["sources"][0]["current_snapshot_id"],
            uploaded["snapshot_id"],
        )

    def test_m3_plan_pack_link_and_coverage_flow(self) -> None:
        uploaded = self.client.post(
            "/api/knowledge/example.com/sources/upload",
            files={"file": ("private-spec.docx", docx_bytes())},
            data={"trust_tier": "hard_fact"},
        ).json()
        source_id = uploaded["source_id"]
        self.client.put(
            f"/api/knowledge/example.com/sources/{source_id}/review",
            json={
                "source_kind": "private_file",
                "trust_tier": "hard_fact",
                "decision": "approve",
                "reason": "Verified product specification.",
            },
        )
        published = self.client.post(
            f"/api/knowledge/example.com/sources/{source_id}/publish"
        )
        self.assertEqual(published.status_code, 200, published.text)

        plan_payload = {
            "retrieval_plan_id": "plan-article-1-v1",
            "article_id": "article-1",
            "outline_version": 1,
            "scopes": [
                {
                    "scope_id": "scope-product-facts",
                    "ordinal": 0,
                    "scope_type": "product_fact",
                    "scope_key": "materials",
                    "title": "Product materials",
                    "query_variants": ["carbon steel wood screw material"],
                    "minimum_hits": 1,
                    "minimum_distinct_sources": 1,
                    "require_hard_fact": True,
                }
            ],
        }
        created_plan = self.client.post(
            "/api/knowledge/example.com/retrieval-plans",
            json=plan_payload,
        )
        self.assertEqual(created_plan.status_code, 201, created_plan.text)
        repeated_plan = self.client.post(
            "/api/knowledge/example.com/retrieval-plans",
            json=plan_payload,
        )
        self.assertEqual(repeated_plan.status_code, 201, repeated_plan.text)

        pack_url = (
            "/api/knowledge/example.com/retrieval-plans/"
            "plan-article-1-v1/scopes/scope-product-facts/evidence-packs"
        )
        built_pack = self.client.post(pack_url, json={"limit": 5})
        self.assertEqual(built_pack.status_code, 201, built_pack.text)
        pack = built_pack.json()
        self.assertEqual(pack["sufficiency"], "sufficient")
        self.assertGreaterEqual(len(pack["hits"]), 1)
        self.assertEqual(
            pack["hits"][0]["provenance"]["snapshot_id"],
            uploaded["snapshot_id"],
        )
        repeated_pack = self.client.post(pack_url, json={"limit": 5})
        self.assertEqual(repeated_pack.status_code, 201, repeated_pack.text)
        self.assertEqual(
            repeated_pack.json()["evidence_pack_id"],
            pack["evidence_pack_id"],
        )

        paragraph_hash = "a" * 64
        link_response = self.client.post(
            "/api/knowledge/example.com/evidence-links",
            json={
                "evidence_link_id": "link-article-1-p1",
                "article_id": "article-1",
                "paragraph_id": "p1",
                "paragraph_hash": paragraph_hash,
                "chunk_id": pack["hits"][0]["chunk_id"],
                "visible_words": 12,
            },
        )
        self.assertEqual(link_response.status_code, 201, link_response.text)

        coverage = self.client.post(
            "/api/knowledge/example.com/articles/article-1/knowledge-coverage",
            json={
                "paragraphs": [
                    {
                        "paragraph_id": "p1",
                        "paragraph_hash": paragraph_hash,
                        "visible_words": 12,
                    }
                ],
                "hard_fact_sentences": [],
            },
        )
        self.assertEqual(coverage.status_code, 200, coverage.text)
        self.assertEqual(coverage.json()["paragraph_coverage"], 1.0)
        self.assertEqual(coverage.json()["hard_fact_coverage"], 1.0)

        stale = self.client.post(
            "/api/knowledge/example.com/articles/article-1/"
            "evidence-links/review-stale",
            json={
                "paragraph_id": "p1",
                "current_paragraph_hash": "b" * 64,
            },
        )
        self.assertEqual(stale.status_code, 200, stale.text)
        self.assertEqual(stale.json()["marked_needs_review"], 1)


if __name__ == "__main__":
    unittest.main()
