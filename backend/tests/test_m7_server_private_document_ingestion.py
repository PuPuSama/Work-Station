from __future__ import annotations

import hashlib
import os
import sys
import unittest
import uuid
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

import sqlalchemy as sa
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from PIL import Image


BACKEND_DIR = Path(__file__).resolve().parents[1]
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from knowledge_agent import (  # noqa: E402
    EMBEDDING_DIMENSIONS,
    ChunkEmbedding,
    DocumentInput,
    KnowledgeAsset,
    KnowledgeChunk,
    KnowledgeProject,
    KnowledgeSource,
    PostgresKnowledgeAssetRepository,
    PostgresKnowledgeRepository,
    SourceSnapshot,
    create_knowledge_engine,
)
from knowledge_agent.library import PostgresKnowledgeLibrary  # noqa: E402
from knowledge_agent.http import router  # noqa: E402
from knowledge_agent.schema import (  # noqa: E402
    knowledge_assets,
    knowledge_chunks,
    knowledge_sources,
    projects,
    snapshot_assets,
    source_snapshots,
)
from server_schema import (  # noqa: E402
    organizations,
    project_memberships,
    project_ownership,
    workspace_users,
)
from services.access_control import (  # noqa: E402
    ActorIdentity,
    PostgresProjectAccessRepository,
    ProjectAccessDenied,
    ProjectAccessService,
)
from services.object_store import (  # noqa: E402
    ObjectStoreError,
    StoredObject,
)
from services.server_auth import (  # noqa: E402
    SERVER_AUTH_COOKIE_NAME,
    ServerActorSessionCodec,
)
from services.server_private_document_ingestion import (  # noqa: E402
    PostgresServerPrivateDocumentIngestion,
    ServerPrivateDocumentUploadConflict,
    ServerPrivateDocumentUploadUnavailable,
)
from services.server_request_security import (  # noqa: E402
    ServerRequestSecurity,
)


DATABASE_URL_ENV = "ARTICLE_AGENT_DATABASE_URL"


def embedded_png_bytes() -> bytes:
    output = BytesIO()
    Image.new("RGB", (12, 8), color=(80, 40, 20)).save(
        output,
        format="PNG",
    )
    return output.getvalue()


def docx_bytes(
    fact: str = "Material: carbon steel. Finish: zinc plated.",
) -> bytes:
    document = Document()
    document.core_properties.title = "Private Fastener Specification"
    document.add_heading("Wood Screw", level=1)
    document.add_paragraph(fact)
    document.add_picture(BytesIO(embedded_png_bytes()))
    output = BytesIO()
    document.save(output)
    return output.getvalue()


class RecordingStore:
    def __init__(self) -> None:
        self.put_calls: list[dict[str, object]] = []
        self.objects: dict[str, bytes] = {}

    def check_ready(self) -> None:
        return None

    def put(
        self,
        *,
        key: str,
        data: bytes,
        content_type: str,
        metadata=None,
    ) -> StoredObject:
        body = bytes(data)
        self.objects[key] = body
        self.put_calls.append(
            {
                "key": key,
                "content_type": content_type,
                "metadata": dict(metadata or {}),
            }
        )
        return StoredObject(
            key=key,
            content_hash=hashlib.sha256(body).hexdigest(),
            content_type=content_type,
            byte_size=len(body),
            etag="test-etag",
        )

    def get(self, key: str, *, max_bytes: int) -> bytes:
        return self.objects[key][: max_bytes + 1]

    def create_download_url(
        self,
        key: str,
        *,
        expires_seconds: int,
    ) -> str:
        del expires_seconds
        return f"https://signed.example.test/{key}"

    def list(self, *, prefix: str):
        del prefix
        return ()

    def delete(self, key: str) -> None:
        self.objects.pop(key, None)


class FailingStore(RecordingStore):
    def put(self, **kwargs) -> StoredObject:
        del kwargs
        raise ObjectStoreError(
            "object storage failed with private-access-key-value"
        )


class RevokingStore(RecordingStore):
    def __init__(
        self,
        engine: sa.Engine,
        *,
        organization_id: str,
        project_id: str,
        user_id: str,
    ) -> None:
        super().__init__()
        self._engine = engine
        self._scope = (organization_id, project_id, user_id)
        self._revoked = False

    def put(self, **kwargs) -> StoredObject:
        stored = super().put(**kwargs)
        if not self._revoked:
            self._revoked = True
            organization_id, project_id, user_id = self._scope
            with self._engine.begin() as connection:
                connection.execute(
                    project_memberships.delete().where(
                        project_memberships.c.organization_id
                        == organization_id,
                        project_memberships.c.project_id == project_id,
                        project_memberships.c.user_id == user_id,
                    )
                )
        return stored


class RecordingAuditWriter:
    def __init__(self) -> None:
        self.events: list[object] = []

    def append(self, connection, event) -> None:
        if not connection.in_transaction():
            raise AssertionError("audit must share the business transaction")
        self.events.append(event)


class FailingAuditWriter:
    def append(self, connection, event) -> None:
        del connection, event
        raise RuntimeError(
            "audit failed with https://private-audit.example/secret"
        )


class CurrentSessionVersions:
    def is_current(self, session) -> bool:
        del session
        return True


@unittest.skipUnless(
    os.environ.get(DATABASE_URL_ENV),
    f"{DATABASE_URL_ENV} is required for PostgreSQL integration tests",
)
class ServerPrivateDocumentIngestionTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.engine = create_knowledge_engine(
            os.environ[DATABASE_URL_ENV]
        )

    @classmethod
    def tearDownClass(cls) -> None:
        cls.engine.dispose()

    def setUp(self) -> None:
        prefix = f"m7-private-upload-{uuid.uuid4().hex}"
        self.organization_id = f"{prefix}-org"
        self.project_id = f"{prefix}.example.test"
        self.other_project_id = f"other-{prefix}.example.test"
        self.editor_id = f"{prefix}-editor"
        self.viewer_id = f"{prefix}-viewer"
        self.repository = PostgresKnowledgeRepository(self.engine)
        for project_id in (self.project_id, self.other_project_id):
            self.repository.upsert_project(
                KnowledgeProject(
                    project_id=project_id,
                    customer_name=project_id,
                    official_domain=project_id,
                )
            )
        with self.engine.begin() as connection:
            connection.execute(
                organizations.insert().values(
                    organization_id=self.organization_id,
                    name="Private Upload Test Organization",
                )
            )
            connection.execute(
                workspace_users.insert(),
                (
                    {
                        "organization_id": self.organization_id,
                        "user_id": self.editor_id,
                        "display_name": "Knowledge Editor",
                    },
                    {
                        "organization_id": self.organization_id,
                        "user_id": self.viewer_id,
                        "display_name": "Knowledge Viewer",
                    },
                ),
            )
            connection.execute(
                project_ownership.insert(),
                (
                    {
                        "organization_id": self.organization_id,
                        "project_id": self.project_id,
                    },
                    {
                        "organization_id": self.organization_id,
                        "project_id": self.other_project_id,
                    },
                ),
            )
            connection.execute(
                project_memberships.insert(),
                (
                    {
                        "organization_id": self.organization_id,
                        "project_id": self.project_id,
                        "user_id": self.editor_id,
                        "role": "editor",
                        "granted_by_user_id": self.editor_id,
                    },
                    {
                        "organization_id": self.organization_id,
                        "project_id": self.project_id,
                        "user_id": self.viewer_id,
                        "role": "viewer",
                        "granted_by_user_id": self.editor_id,
                    },
                ),
            )
        self.editor = ActorIdentity(
            self.organization_id,
            self.editor_id,
        )
        self.viewer = ActorIdentity(
            self.organization_id,
            self.viewer_id,
        )
        # A DOCX ZIP carries generation timestamps. Reuse one immutable body
        # per test so an idempotency assertion never crosses a clock boundary
        # and accidentally submits different bytes.
        self.document_content = docx_bytes()
        self.store = RecordingStore()
        self.audit = RecordingAuditWriter()
        self.service = self._service(
            store=self.store,
            audit=self.audit,
        )

    def tearDown(self) -> None:
        project_ids = (self.project_id, self.other_project_id)
        with self.engine.begin() as connection:
            connection.execute(
                snapshot_assets.delete().where(
                    snapshot_assets.c.project_id.in_(project_ids)
                )
            )
            connection.execute(
                knowledge_assets.delete().where(
                    knowledge_assets.c.project_id.in_(project_ids)
                )
            )
            connection.execute(
                knowledge_chunks.delete().where(
                    knowledge_chunks.c.project_id.in_(project_ids)
                )
            )
            connection.execute(
                source_snapshots.delete().where(
                    source_snapshots.c.project_id.in_(project_ids)
                )
            )
            connection.execute(
                knowledge_sources.delete().where(
                    knowledge_sources.c.project_id.in_(project_ids)
                )
            )
            connection.execute(
                project_memberships.delete().where(
                    project_memberships.c.organization_id
                    == self.organization_id
                )
            )
            connection.execute(
                project_ownership.delete().where(
                    project_ownership.c.organization_id
                    == self.organization_id
                )
            )
            connection.execute(
                workspace_users.delete().where(
                    workspace_users.c.organization_id
                    == self.organization_id
                )
            )
            connection.execute(
                organizations.delete().where(
                    organizations.c.organization_id
                    == self.organization_id
                )
            )
            connection.execute(
                projects.delete().where(
                    projects.c.project_id.in_(project_ids)
                )
            )

    def _service(
        self,
        *,
        store,
        audit,
    ) -> PostgresServerPrivateDocumentIngestion:
        return PostgresServerPrivateDocumentIngestion(
            self.engine,
            store=store,
            bucket="private-test-bucket",
            audit=audit,
        )

    def _upload(
        self,
        service: PostgresServerPrivateDocumentIngestion | None = None,
        *,
        actor: ActorIdentity | None = None,
        source_id: str = "private-spec",
        display_name: str = "Private specification",
        content: bytes | None = None,
    ):
        return (service or self.service).upload(
            actor=actor or self.editor,
            project_id=self.project_id,
            source_id=source_id,
            display_name=display_name,
            document_input=DocumentInput(
                filename="private-spec.docx",
                content=(
                    self.document_content if content is None else content
                ),
                content_type=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
            ),
            trust_tier="hard_fact",
        )

    def _seed_published_source(self, source_id: str) -> SourceSnapshot:
        self.repository.upsert_source(
            KnowledgeSource(
                project_id=self.project_id,
                source_id=source_id,
                display_name="Published private specification",
                source_kind="private_file",
                trust_tier="hard_fact",
                status="inbox",
            )
        )
        snapshot = SourceSnapshot(
            project_id=self.project_id,
            source_id=source_id,
            snapshot_id=f"{source_id}-current",
            content_hash="a" * 64,
            fetched_at=datetime(2026, 7, 31, tzinfo=timezone.utc),
            parser_name="test",
            parser_version="1",
        )
        chunk = KnowledgeChunk(
            project_id=self.project_id,
            source_id=source_id,
            snapshot_id=snapshot.snapshot_id,
            chunk_id=f"{snapshot.snapshot_id}:000000",
            text="Old published private evidence.",
        )
        self.repository.store_snapshot(
            self.project_id,
            snapshot,
            (chunk,),
        )
        self.repository.store_embeddings(
            self.project_id,
            (
                ChunkEmbedding(
                    project_id=self.project_id,
                    chunk_id=chunk.chunk_id,
                    snapshot_id=snapshot.snapshot_id,
                    vector=(1.0,)
                    + (0.0,) * (EMBEDDING_DIMENSIONS - 1),
                    embedding_model="published-test-model",
                ),
            ),
        )
        self.repository.activate_snapshot(
            self.project_id,
            source_id,
            snapshot.snapshot_id,
            "published-test-model",
        )
        return snapshot

    def _row_counts(self) -> dict[str, int]:
        with self.engine.connect() as connection:
            return {
                "sources": connection.execute(
                    sa.select(sa.func.count())
                    .select_from(knowledge_sources)
                    .where(
                        knowledge_sources.c.project_id == self.project_id
                    )
                ).scalar_one(),
                "snapshots": connection.execute(
                    sa.select(sa.func.count())
                    .select_from(source_snapshots)
                    .where(
                        source_snapshots.c.project_id == self.project_id
                    )
                ).scalar_one(),
                "chunks": connection.execute(
                    sa.select(sa.func.count())
                    .select_from(knowledge_chunks)
                    .where(
                        knowledge_chunks.c.project_id == self.project_id
                    )
                ).scalar_one(),
            }

    def test_upload_is_project_scoped_atomic_and_retry_idempotent(self) -> None:
        first = self._upload()
        second = self._upload(display_name="Ignored retry title")

        self.assertTrue(first.created)
        self.assertFalse(second.created)
        self.assertEqual(
            first.result.snapshot.snapshot_id,
            second.result.snapshot.snapshot_id,
        )
        self.assertEqual(
            first.result.source.pending_snapshot_id,
            first.result.snapshot.snapshot_id,
        )
        self.assertEqual(
            second.result.source.pending_snapshot_id,
            first.result.snapshot.snapshot_id,
        )
        counts = self._row_counts()
        self.assertEqual(counts["sources"], 1)
        self.assertEqual(counts["snapshots"], 1)
        self.assertGreater(counts["chunks"], 0)
        stored_source = PostgresKnowledgeLibrary(self.engine).get_source(
            self.project_id,
            "private-spec",
        )
        self.assertEqual(
            stored_source.display_name,
            "Private specification",
        )
        self.assertIsNone(stored_source.current_snapshot_id)
        self.assertEqual(
            stored_source.pending_snapshot_id,
            first.result.snapshot.snapshot_id,
        )
        self.assertEqual(
            [event.action for event in self.audit.events],
            ["knowledge.source.uploaded"],
        )
        event = self.audit.events[0]
        self.assertEqual(event.project_id, self.project_id)
        self.assertEqual(event.target_id, "private-spec")
        self.assertEqual(
            set(event.details),
            {
                "snapshot_id",
                "parser_name",
                "parser_version",
                "chunk_count",
                "asset_count",
            },
        )
        self.assertNotIn("filename", str(event.details).casefold())
        self.assertNotIn("s3://", str(event.details).casefold())
        expected_prefix = (
            f"organizations/{self.organization_id}/"
            f"projects/{self.project_id}/blobs/"
        )
        self.assertTrue(self.store.put_calls)
        self.assertTrue(
            all(
                str(call["key"]).startswith(expected_prefix)
                for call in self.store.put_calls
            )
        )

    def test_viewer_is_rejected_before_any_object_or_database_write(
        self,
    ) -> None:
        with self.assertRaisesRegex(
            ProjectAccessDenied,
            "^project access denied$",
        ):
            self._upload(actor=self.viewer)

        self.assertEqual(self.store.put_calls, [])
        self.assertEqual(
            self._row_counts(),
            {"sources": 0, "snapshots": 0, "chunks": 0},
        )
        self.assertEqual(self.audit.events, [])

    def test_editor_cannot_upload_to_an_unassigned_project(self) -> None:
        with self.assertRaisesRegex(
            ProjectAccessDenied,
            "^project access denied$",
        ):
            self.service.upload(
                actor=self.editor,
                project_id=self.other_project_id,
                source_id="cross-project-private-spec",
                display_name="Cross-project private specification",
                document_input=DocumentInput(
                    filename="private-spec.docx",
                    content=docx_bytes(),
                ),
            )

        self.assertEqual(self.store.put_calls, [])
        with self.engine.connect() as connection:
            source_count = connection.execute(
                sa.select(sa.func.count())
                .select_from(knowledge_sources)
                .where(
                    knowledge_sources.c.project_id
                    == self.other_project_id
                )
            ).scalar_one()
        self.assertEqual(source_count, 0)
        self.assertEqual(self.audit.events, [])

    def test_existing_deduplicated_asset_id_is_used_by_snapshot_link(
        self,
    ) -> None:
        content_hash = hashlib.sha256(embedded_png_bytes()).hexdigest()
        PostgresKnowledgeAssetRepository(self.engine).put_asset(
            KnowledgeAsset(
                project_id=self.project_id,
                asset_id="legacy-product-image",
                content_hash=content_hash,
                artifact_uri=(
                    "s3://private-test-bucket/"
                    f"organizations/{self.organization_id}/"
                    f"projects/{self.project_id}/blobs/"
                    f"{content_hash[:2]}/{content_hash}"
                ),
                content_type="image/png",
                byte_size=len(embedded_png_bytes()),
                width=12,
                height=8,
            )
        )

        uploaded = self._upload()

        self.assertEqual(
            [asset.asset_id for asset in uploaded.result.assets],
            ["legacy-product-image"],
        )
        self.assertEqual(
            [
                link.asset_id
                for link in uploaded.result.snapshot_assets
            ],
            ["legacy-product-image"],
        )
        with self.engine.connect() as connection:
            linked_id = connection.execute(
                sa.select(snapshot_assets.c.asset_id).where(
                    snapshot_assets.c.project_id == self.project_id,
                    snapshot_assets.c.source_id == "private-spec",
                )
            ).scalar_one()
        self.assertEqual(linked_id, "legacy-product-image")

    def test_published_source_accepts_one_pending_snapshot_and_retry(
        self,
    ) -> None:
        source_id = "published-private-spec"
        old_snapshot = self._seed_published_source(source_id)

        first = self._upload(
            source_id=source_id,
            display_name="New private specification",
        )
        first_counts = self._row_counts()
        retry = self._upload(
            source_id=source_id,
            display_name="Ignored retry title",
        )

        self.assertTrue(first.created)
        self.assertFalse(retry.created)
        self.assertEqual(
            retry.result.snapshot.snapshot_id,
            first.result.snapshot.snapshot_id,
        )
        self.assertEqual(self._row_counts(), first_counts)
        self.assertEqual(
            [event.action for event in self.audit.events],
            ["knowledge.source.uploaded"],
        )
        source = PostgresKnowledgeLibrary(self.engine).get_source(
            self.project_id,
            source_id,
        )
        self.assertEqual(source.status, "published")
        self.assertEqual(
            source.current_snapshot_id,
            old_snapshot.snapshot_id,
        )
        self.assertEqual(
            source.pending_snapshot_id,
            first.result.snapshot.snapshot_id,
        )
        self.assertEqual(
            source.display_name,
            "Published private specification",
        )
        self.assertEqual(first_counts["snapshots"], 2)

        different_content = docx_bytes(
            "Material: stainless steel. Finish: black oxide."
        )
        with self.assertRaisesRegex(
            ServerPrivateDocumentUploadConflict,
            "^knowledge source already has a pending snapshot$",
        ):
            self._upload(
                source_id=source_id,
                display_name="Second pending specification",
                content=different_content,
            )

        unchanged = PostgresKnowledgeLibrary(self.engine).get_source(
            self.project_id,
            source_id,
        )
        self.assertEqual(unchanged.status, "published")
        self.assertEqual(
            unchanged.current_snapshot_id,
            old_snapshot.snapshot_id,
        )
        self.assertEqual(
            unchanged.pending_snapshot_id,
            first.result.snapshot.snapshot_id,
        )
        self.assertEqual(self._row_counts(), first_counts)
        self.assertEqual(len(self.audit.events), 1)

    def test_deduplication_rejects_assets_outside_server_scope(
        self,
    ) -> None:
        content_hash = hashlib.sha256(embedded_png_bytes()).hexdigest()
        invalid_uris = (
            f"file:///legacy/{content_hash}.png",
            (
                "s3://private-test-bucket/organizations/other-org/"
                f"projects/{self.project_id}/blobs/"
                f"{content_hash[:2]}/{content_hash}"
            ),
            (
                "s3://other-bucket/"
                f"organizations/{self.organization_id}/"
                f"projects/{self.project_id}/blobs/"
                f"{content_hash[:2]}/{content_hash}"
            ),
            (
                "s3://private-test-bucket/"
                f"organizations/{self.organization_id}/"
                f"projects/{self.project_id}/blobs/"
                f"{content_hash[:2]}/not-the-content-hash"
            ),
        )
        for index, artifact_uri in enumerate(invalid_uris):
            with self.subTest(artifact_uri=artifact_uri):
                with self.engine.begin() as connection:
                    connection.execute(
                        knowledge_assets.delete().where(
                            knowledge_assets.c.project_id
                            == self.project_id,
                            knowledge_assets.c.content_hash
                            == content_hash,
                        )
                    )
                PostgresKnowledgeAssetRepository(self.engine).put_asset(
                    KnowledgeAsset(
                        project_id=self.project_id,
                        asset_id=f"legacy-image-{index}",
                        content_hash=content_hash,
                        artifact_uri=artifact_uri,
                        content_type="image/png",
                        byte_size=len(embedded_png_bytes()),
                        width=12,
                        height=8,
                    )
                )

                with self.assertRaisesRegex(
                    ServerPrivateDocumentUploadConflict,
                    "^deduplicated asset is outside the server scope$",
                ):
                    self._upload(
                        source_id=f"invalid-asset-source-{index}"
                    )

                self.assertEqual(
                    self._row_counts(),
                    {"sources": 0, "snapshots": 0, "chunks": 0},
                )
                self.assertEqual(self.audit.events, [])

    def test_permission_revocation_after_object_write_rolls_back_database(
        self,
    ) -> None:
        source_id = "revoked-pending-private-spec"
        current = self._seed_published_source(source_id)
        baseline = self._row_counts()
        store = RevokingStore(
            self.engine,
            organization_id=self.organization_id,
            project_id=self.project_id,
            user_id=self.editor_id,
        )
        service = self._service(store=store, audit=self.audit)

        with self.assertRaisesRegex(
            ProjectAccessDenied,
            "^project access denied$",
        ):
            self._upload(service, source_id=source_id)

        self.assertTrue(store.put_calls)
        self.assertEqual(self._row_counts(), baseline)
        source = PostgresKnowledgeLibrary(self.engine).get_source(
            self.project_id,
            source_id,
        )
        self.assertEqual(source.status, "published")
        self.assertEqual(source.current_snapshot_id, current.snapshot_id)
        self.assertIsNone(source.pending_snapshot_id)
        self.assertEqual(self.audit.events, [])

    def test_audit_failure_rolls_back_all_database_rows(self) -> None:
        source_id = "audit-failed-pending-private-spec"
        current = self._seed_published_source(source_id)
        baseline = self._row_counts()
        service = self._service(
            store=self.store,
            audit=FailingAuditWriter(),
        )

        with self.assertRaisesRegex(
            ServerPrivateDocumentUploadUnavailable,
            "^private document could not be committed$",
        ) as captured:
            self._upload(service, source_id=source_id)

        self.assertNotIn("private-audit", str(captured.exception))
        self.assertTrue(self.store.put_calls)
        self.assertEqual(self._row_counts(), baseline)
        source = PostgresKnowledgeLibrary(self.engine).get_source(
            self.project_id,
            source_id,
        )
        self.assertEqual(source.status, "published")
        self.assertEqual(source.current_snapshot_id, current.snapshot_id)
        self.assertIsNone(source.pending_snapshot_id)

    def _http_client(
        self,
        *,
        service: PostgresServerPrivateDocumentIngestion,
    ) -> tuple[TestClient, ServerActorSessionCodec]:
        codec = ServerActorSessionCodec(b"k" * 32)
        app = FastAPI()
        app.state.server_mode_enabled = True
        app.state.knowledge_agent_runtime = SimpleNamespace()
        app.state.server_private_document_ingestion = service
        app.state.server_request_security = ServerRequestSecurity(
            codec=codec,
            access=ProjectAccessService(
                PostgresProjectAccessRepository(self.engine)
            ),
            sessions=CurrentSessionVersions(),
        )
        app.include_router(router)
        return TestClient(app), codec

    def test_server_http_upload_reports_created_and_exact_retry(self) -> None:
        client, codec = self._http_client(service=self.service)
        client.cookies.set(
            SERVER_AUTH_COOKIE_NAME,
            codec.create(self.editor),
        )
        content = docx_bytes()
        try:
            first = client.post(
                f"/api/knowledge/{self.project_id}/sources/upload",
                files={"file": ("private-spec.docx", content)},
                data={
                    "source_id": "http-private-spec",
                    "display_name": "HTTP private specification",
                    "trust_tier": "hard_fact",
                },
            )
            second = client.post(
                f"/api/knowledge/{self.project_id}/sources/upload",
                files={"file": ("private-spec.docx", content)},
                data={
                    "source_id": "http-private-spec",
                    "display_name": "HTTP private specification",
                    "trust_tier": "hard_fact",
                },
            )
        finally:
            client.close()

        self.assertEqual(first.status_code, 201, first.text)
        self.assertEqual(second.status_code, 201, second.text)
        self.assertTrue(first.json()["created"])
        self.assertFalse(second.json()["created"])
        self.assertEqual(
            first.json()["snapshot_id"],
            second.json()["snapshot_id"],
        )
        self.assertNotIn("s3://", first.text)
        self.assertNotIn("content_hash", first.text)

    def test_server_http_hides_object_store_failure_details(self) -> None:
        service = self._service(
            store=FailingStore(),
            audit=self.audit,
        )
        client, codec = self._http_client(service=service)
        client.cookies.set(
            SERVER_AUTH_COOKIE_NAME,
            codec.create(self.editor),
        )
        try:
            response = client.post(
                f"/api/knowledge/{self.project_id}/sources/upload",
                files={"file": ("private-spec.docx", docx_bytes())},
            )
        finally:
            client.close()

        self.assertEqual(response.status_code, 503, response.text)
        self.assertEqual(
            response.json()["detail"],
            "Private document ingestion is temporarily unavailable.",
        )
        self.assertNotIn("private-access-key-value", response.text)


if __name__ == "__main__":
    unittest.main()
