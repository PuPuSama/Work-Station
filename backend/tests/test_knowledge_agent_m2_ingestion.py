from __future__ import annotations

import sys
import tempfile
import unittest
from hashlib import sha256
from io import BytesIO
from pathlib import Path

from docx import Document
from PIL import Image


BACKEND_DIR = Path(__file__).resolve().parents[1]
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from knowledge_agent import (  # noqa: E402
    DocumentInput,
    KnowledgeAsset,
    KnowledgeSource,
    LocalKnowledgeArtifactStore,
    ParsedDocumentChunker,
    PrivateDocumentIngestionService,
    SnapshotAsset,
)


class FakeKnowledgeRepository:
    def __init__(self) -> None:
        self.sources: list[KnowledgeSource] = []
        self.snapshots: list[tuple[str, object, tuple[object, ...]]] = []

    def upsert_project(self, project: object) -> None:
        raise NotImplementedError

    def upsert_source(self, source: KnowledgeSource) -> None:
        self.sources.append(source)

    def store_snapshot(
        self, project_id: str, snapshot: object, chunks: tuple[object, ...]
    ) -> None:
        self.snapshots.append((project_id, snapshot, tuple(chunks)))

    def store_embeddings(self, project_id: str, embeddings: object) -> None:
        raise NotImplementedError

    def activate_snapshot(
        self,
        project_id: str,
        source_id: str,
        snapshot_id: str,
        embedding_model: str,
    ) -> None:
        raise NotImplementedError

    def get_chunks(self, project_id: str, chunk_ids: object) -> tuple[object, ...]:
        raise NotImplementedError


class FakeAssetRepository:
    def __init__(self) -> None:
        self.assets: dict[str, KnowledgeAsset] = {}
        self.links: list[SnapshotAsset] = []

    def put_asset(self, asset: KnowledgeAsset) -> KnowledgeAsset:
        existing = next(
            (
                item
                for item in self.assets.values()
                if item.project_id == asset.project_id
                and item.content_hash == asset.content_hash
            ),
            None,
        )
        if existing is not None:
            return existing
        self.assets[asset.asset_id] = asset
        return asset

    def link_snapshot_asset(self, link: SnapshotAsset) -> None:
        if link not in self.links:
            self.links.append(link)

    def get_asset(self, project_id: str, asset_id: str) -> KnowledgeAsset | None:
        return self.assets.get(asset_id)

    def list_snapshot_assets(
        self, project_id: str, source_id: str, snapshot_id: str
    ) -> tuple[SnapshotAsset, ...]:
        return tuple(
            link
            for link in self.links
            if link.project_id == project_id
            and link.source_id == source_id
            and link.snapshot_id == snapshot_id
        )


def document_bytes() -> bytes:
    document = Document()
    document.add_heading("Fastener Specification", level=1)
    document.add_paragraph("A" * 420)
    image = BytesIO()
    Image.new("RGB", (10, 10), color=(80, 40, 20)).save(image, format="PNG")
    image.seek(0)
    document.add_picture(image)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


class ParsedDocumentChunkerTests(unittest.TestCase):
    def test_long_blocks_split_with_stable_snapshot_prefixed_ids(self) -> None:
        from knowledge_agent import DocumentParserRouter

        parsed = DocumentParserRouter().parse(
            DocumentInput(filename="spec.docx", content=document_bytes())
        )
        chunker = ParsedDocumentChunker(max_characters=256)

        first = chunker.chunk(
            project_id="project-a",
            source_id="source-a",
            snapshot_id="snapshot-a",
            document=parsed,
        )
        second = chunker.chunk(
            project_id="project-a",
            source_id="source-a",
            snapshot_id="snapshot-a",
            document=parsed,
        )

        self.assertEqual(first, second)
        self.assertGreaterEqual(len(first), 3)
        self.assertEqual(first[0].chunk_id, "snapshot-a:000000")
        self.assertTrue(all(len(chunk.text) <= 256 for chunk in first))
        self.assertEqual(dict(first[-1].locator)["part_count"], 2)


class LocalArtifactStoreTests(unittest.TestCase):
    def test_store_is_content_addressed_idempotent_and_project_scoped(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            store = LocalKnowledgeArtifactStore(Path(directory))
            digest = sha256(b"same").hexdigest()
            first = store.put(
                project_id="project-a",
                namespace="raw",
                content_hash=digest,
                filename="spec.docx",
                content=b"same",
            )
            second = store.put(
                project_id="project-a",
                namespace="raw",
                content_hash=digest,
                filename="spec.docx",
                content=b"same",
            )

            self.assertEqual(first, second)
            self.assertEqual(Path(first.removeprefix("file:///")).name, "spec.docx")
            with self.assertRaisesRegex(ValueError, "does not match"):
                store.put(
                    project_id="project-a",
                    namespace="raw",
                    content_hash="a" * 64,
                    filename="spec.docx",
                    content=b"same",
                )
            with self.assertRaisesRegex(ValueError, "project_id"):
                store.put(
                    project_id="../foreign",
                    namespace="raw",
                    content_hash=digest,
                    filename="spec.docx",
                    content=b"same",
                )


class PrivateDocumentIngestionServiceTests(unittest.TestCase):
    def test_ingestion_is_stable_reviewable_and_persists_asset_evidence(self) -> None:
        knowledge_repository = FakeKnowledgeRepository()
        asset_repository = FakeAssetRepository()
        with tempfile.TemporaryDirectory() as directory:
            service = PrivateDocumentIngestionService(
                repository=knowledge_repository,  # type: ignore[arg-type]
                asset_repository=asset_repository,
                artifact_store=LocalKnowledgeArtifactStore(Path(directory)),
                chunker=ParsedDocumentChunker(max_characters=256),
            )
            source = DocumentInput(filename="spec.docx", content=document_bytes())
            first = service.ingest(
                project_id="project-a",
                source_id="private-spec",
                display_name="Private specification",
                document_input=source,
            )
            second = service.ingest(
                project_id="project-a",
                source_id="private-spec",
                display_name="Private specification",
                document_input=source,
            )

        self.assertEqual(first.snapshot.snapshot_id, second.snapshot.snapshot_id)
        self.assertEqual(first.chunks, second.chunks)
        self.assertEqual(first.source.status, "inbox")
        self.assertIsNone(first.source.current_snapshot_id)
        self.assertTrue(first.snapshot.raw_artifact_uri)
        self.assertTrue(first.snapshot.normalized_artifact_uri)
        self.assertEqual(len(first.assets), 1)
        self.assertEqual(len(first.snapshot_assets), 1)
        self.assertEqual(
            first.snapshot_assets[0].snapshot_id,
            first.snapshot.snapshot_id,
        )
        self.assertEqual(len(asset_repository.assets), 1)
        self.assertEqual(len(asset_repository.links), 1)
        self.assertNotEqual(
            first.snapshot.raw_artifact_uri,
            first.snapshot.normalized_artifact_uri,
        )

if __name__ == "__main__":
    unittest.main()
