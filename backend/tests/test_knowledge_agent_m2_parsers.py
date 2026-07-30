from __future__ import annotations

import sys
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from PIL import Image
from pypdf import PdfWriter


BACKEND_DIR = Path(__file__).resolve().parents[1]
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from knowledge_agent import (  # noqa: E402
    DocumentInput,
    DocumentParseError,
    DocumentParserRouter,
    ParsedBlock,
    ParsedDocument,
    UnsupportedDocumentError,
)


def docx_bytes(*, with_image: bool = False) -> bytes:
    document = Document()
    document.core_properties.title = "Fastener data sheet"
    document.add_heading("Wood Screws", level=1)
    document.add_paragraph("Carbon steel screw for timber applications.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Diameter"
    table.cell(0, 1).text = "Length"
    table.cell(1, 0).text = "4.0 mm"
    table.cell(1, 1).text = "50 mm"
    if with_image:
        image_stream = BytesIO()
        Image.new("RGB", (8, 8), color=(30, 80, 120)).save(
            image_stream, format="PNG"
        )
        image_stream.seek(0)
        document.add_picture(image_stream)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def xlsx_bytes() -> bytes:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Products"
    worksheet.append(("SKU", "Material", "Active"))
    worksheet.append(("WS-001", "Carbon Steel", True))
    output = BytesIO()
    workbook.save(output)
    workbook.close()
    return output.getvalue()


class DocumentInputContractTests(unittest.TestCase):
    def test_input_computes_hash_and_rejects_paths(self) -> None:
        document = DocumentInput(
            filename="catalog.docx",
            content=b"document bytes",
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document; charset=binary"
            ),
        )

        self.assertEqual(len(document.content_hash), 64)
        self.assertEqual(document.suffix, ".docx")
        self.assertNotIn("charset", document.content_type or "")
        with self.assertRaisesRegex(ValueError, "directory path"):
            DocumentInput(filename="../catalog.docx", content=b"x")
        with self.assertRaisesRegex(ValueError, "non-empty bytes"):
            DocumentInput(filename="empty.docx", content=b"")

    def test_parsed_document_requires_stable_ordinals(self) -> None:
        block = ParsedBlock(kind="paragraph", ordinal=1, text="Evidence")
        with self.assertRaisesRegex(ValueError, "contiguous"):
            ParsedDocument(
                filename="evidence.docx",
                content_type="application/test",
                content_hash="a" * 64,
                parser_name="fake",
                parser_version="1",
                blocks=(block,),
            )


class DocumentParserRouterTests(unittest.TestCase):
    def setUp(self) -> None:
        self.router = DocumentParserRouter()

    def test_routes_docx_and_preserves_heading_table_and_image_evidence(self) -> None:
        source = DocumentInput(
            filename="fasteners.docx",
            content=docx_bytes(with_image=True),
        )

        parsed = self.router.parse(source)

        self.assertEqual(parsed.parser_name, "docx-lightweight")
        self.assertEqual(parsed.title, "Fastener data sheet")
        self.assertEqual(parsed.content_hash, source.content_hash)
        self.assertEqual(parsed.blocks[0].kind, "heading")
        self.assertEqual(parsed.blocks[0].heading_path, ("Wood Screws",))
        self.assertTrue(
            any(
                block.kind == "table_row" and "Diameter\tLength" in block.text
                for block in parsed.blocks
            )
        )
        self.assertEqual(len(parsed.assets), 1)
        self.assertEqual(parsed.assets[0].content_type, "image/png")
        self.assertEqual(len(parsed.assets[0].content_hash), 64)

    def test_routes_excel_with_sheet_and_row_locators(self) -> None:
        parsed = self.router.parse(
            DocumentInput(filename="catalog.xlsx", content=xlsx_bytes())
        )

        self.assertEqual(parsed.parser_name, "openpyxl-lightweight")
        self.assertEqual(len(parsed.blocks), 2)
        self.assertEqual(parsed.blocks[0].text, "SKU\tMaterial\tActive")
        self.assertEqual(parsed.blocks[1].text, "WS-001\tCarbon Steel\ttrue")
        self.assertEqual(parsed.blocks[1].heading_path, ("Products",))
        self.assertEqual(
            dict(parsed.blocks[1].locator),
            {"sheet": "Products", "row_number": 2},
        )

    def test_rejects_unsupported_and_conflicting_metadata(self) -> None:
        with self.assertRaisesRegex(UnsupportedDocumentError, "no document parser"):
            self.router.parse(DocumentInput(filename="notes.txt", content=b"hello"))

        with self.assertRaisesRegex(DocumentParseError, "PDF document"):
            writer = PdfWriter()
            writer.add_blank_page(width=72, height=72)
            output = BytesIO()
            writer.write(output)
            self.router.parse(DocumentInput(filename="blank.pdf", content=output.getvalue()))

        with self.assertRaisesRegex(
            ValueError, "multiple parsers"
        ):
            self.router.parse(
                DocumentInput(
                    filename="wrong.docx",
                    content=docx_bytes(),
                    content_type="application/pdf",
                )
            )


if __name__ == "__main__":
    unittest.main()
