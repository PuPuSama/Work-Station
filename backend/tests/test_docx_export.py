from __future__ import annotations

import sys
import tempfile
import unittest
import zipfile
from pathlib import Path
from types import SimpleNamespace
from xml.etree import ElementTree

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image


BACKEND_DIR = Path(__file__).resolve().parents[1]
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from services.article_images import (  # noqa: E402
    ArticleImageError,
    ImageAnchorRequiredError,
    prepare_task_images,
)
from services.docx_export import export_task_docx  # noqa: E402


WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"

FAQ_BLOCK = """## FAQ

**Q: What should buyers check first?**

A: Buyers should check the application requirements first.

**Q: When should buyers request a sample?**

A: Buyers should request one before approval when fit matters.

**Q: Why should buyers compare suppliers?**

A: Buyers should compare capability, quality control, delivery, and support.
"""


def make_png(path: Path, color: tuple[int, int, int]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    Image.new("RGB", (180, 108), color).save(path, format="PNG")


def test_config() -> SimpleNamespace:
    return SimpleNamespace(
        docx_font="Arial",
        title_1_size=20.0,
        title_2_size=16.0,
        title_3_size=13.0,
        body_size=11.0,
    )


def paragraph_text(paragraph) -> str:
    return "".join(node.text or "" for node in paragraph._p.iter() if node.tag == f"{{{WORD_NS}}}t")


class DocxExportTests(unittest.TestCase):
    def test_homepage_link_uses_project_brand_name_in_export(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            task_dir = Path(temporary) / "task"
            article = """# Brand Link Export

This transition introduces the [Company Website](https://company.example/).

## First Section

This section gives buyers practical detail.

""" + FAQ_BLOCK
            task = SimpleNamespace(
                task_dir=str(task_dir),
                customer="https://company.example/",
                brand_name="Acme Fasteners",
                selected_title="Brand Link Export",
                topic="Topic",
                article=article,
                final_article="",
                linked_article="",
                humanized_article="",
                initial_article="",
                hero_image="",
                products=[],
                images=[],
            )

            output = export_task_docx(test_config(), task)
            exported = Document(output)
            display = [paragraph_text(paragraph) for paragraph in exported.paragraphs]
            self.assertIn(
                "This transition introduces the Acme Fasteners.",
                display,
            )
            self.assertNotIn("Company Website", "\n".join(display))

            with zipfile.ZipFile(output) as archive:
                document_xml = ElementTree.fromstring(archive.read("word/document.xml"))
                hyperlinks = document_xml.findall(f".//{{{WORD_NS}}}hyperlink")
                self.assertEqual(len(hyperlinks), 1)
                hyperlink = hyperlinks[0]
                self.assertEqual(
                    "".join(
                        node.text or ""
                        for node in hyperlink.findall(f".//{{{WORD_NS}}}t")
                    ),
                    "Acme Fasteners",
                )
                self.assertEqual(
                    hyperlink.find(f".//{{{WORD_NS}}}color").attrib[f"{{{WORD_NS}}}val"],
                    "0563C1",
                )

    def test_markdown_table_exports_as_a_real_styled_docx_table(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            task_dir = Path(temporary) / "task"
            article = """# Product Comparison

This transition introduces the product comparison.

## Comparison Table

| Product | Quantity | Product Page |
| :--- | ---: | :---: |
| **Product Alpha** | 12 | [View Product](https://example.com/product-alpha) |
| Special \\| Product | 8 | **[Open Product](https://example.com/product-special)** |
Plain Product | 5 | Ends with an escaped pipe \\|

The table helps buyers compare the available options.

""" + FAQ_BLOCK
            task = SimpleNamespace(
                task_dir=str(task_dir),
                selected_title="Product Comparison",
                topic="Topic",
                article=article,
                final_article="",
                linked_article="",
                humanized_article="",
                initial_article="",
                hero_image="",
                products=[],
                images=[],
            )

            output = export_task_docx(test_config(), task)
            exported = Document(output)
            self.assertEqual(len(exported.tables), 1)
            table = exported.tables[0]
            self.assertEqual(len(table.rows), 4)
            self.assertEqual(len(table.columns), 3)
            visible_cells = [
                [paragraph_text(cell.paragraphs[0]) for cell in row.cells]
                for row in table.rows
            ]
            self.assertEqual(
                visible_cells,
                [
                    ["Product", "Quantity", "Product Page"],
                    ["Product Alpha", "12", "View Product"],
                    ["Special | Product", "8", "Open Product"],
                    ["Plain Product", "5", "Ends with an escaped pipe |"],
                ],
            )
            self.assertEqual(
                [cell.paragraphs[0].alignment for cell in table.rows[1].cells],
                [
                    WD_ALIGN_PARAGRAPH.LEFT,
                    WD_ALIGN_PARAGRAPH.RIGHT,
                    WD_ALIGN_PARAGRAPH.CENTER,
                ],
            )

            with zipfile.ZipFile(output) as archive:
                document_xml = ElementTree.fromstring(archive.read("word/document.xml"))
                relationships = ElementTree.fromstring(
                    archive.read("word/_rels/document.xml.rels")
                )
                table_xml = document_xml.find(f".//{{{WORD_NS}}}tbl")
                self.assertIsNotNone(table_xml)

                table_width = table_xml.find(
                    f"{{{WORD_NS}}}tblPr/{{{WORD_NS}}}tblW"
                )
                table_indent = table_xml.find(
                    f"{{{WORD_NS}}}tblPr/{{{WORD_NS}}}tblInd"
                )
                table_layout = table_xml.find(
                    f"{{{WORD_NS}}}tblPr/{{{WORD_NS}}}tblLayout"
                )
                width_value = int(table_width.attrib[f"{{{WORD_NS}}}w"])
                self.assertEqual(table_width.attrib[f"{{{WORD_NS}}}type"], "dxa")
                self.assertEqual(table_indent.attrib[f"{{{WORD_NS}}}w"], "120")
                self.assertEqual(table_layout.attrib[f"{{{WORD_NS}}}type"], "fixed")

                grid_widths = [
                    int(column.attrib[f"{{{WORD_NS}}}w"])
                    for column in table_xml.findall(
                        f"{{{WORD_NS}}}tblGrid/{{{WORD_NS}}}gridCol"
                    )
                ]
                self.assertEqual(sum(grid_widths), width_value)
                for row_xml in table_xml.findall(f"{{{WORD_NS}}}tr"):
                    cell_widths = [
                        int(
                            cell.find(
                                f"{{{WORD_NS}}}tcPr/{{{WORD_NS}}}tcW"
                            ).attrib[f"{{{WORD_NS}}}w"]
                        )
                        for cell in row_xml.findall(f"{{{WORD_NS}}}tc")
                    ]
                    self.assertEqual(cell_widths, grid_widths)

                header_row = table_xml.find(f"{{{WORD_NS}}}tr")
                self.assertIsNotNone(
                    header_row.find(
                        f"{{{WORD_NS}}}trPr/{{{WORD_NS}}}tblHeader"
                    )
                )
                header_fills = {
                    shading.attrib.get(f"{{{WORD_NS}}}fill")
                    for shading in header_row.findall(
                        f".//{{{WORD_NS}}}tcPr/{{{WORD_NS}}}shd"
                    )
                }
                self.assertEqual(header_fills, {"E7E6E6"})

                for run in table_xml.findall(f".//{{{WORD_NS}}}r"):
                    fonts = run.find(f"{{{WORD_NS}}}rPr/{{{WORD_NS}}}rFonts")
                    self.assertIsNotNone(fonts)
                    self.assertEqual(
                        fonts.attrib[f"{{{WORD_NS}}}ascii"],
                        "Times New Roman",
                    )

                hyperlinks = table_xml.findall(f".//{{{WORD_NS}}}hyperlink")
                self.assertEqual(len(hyperlinks), 2)
                linked_text = {
                    "".join(
                        node.text or ""
                        for node in hyperlink.findall(f".//{{{WORD_NS}}}t")
                    )
                    for hyperlink in hyperlinks
                }
                self.assertEqual(linked_text, {"View Product", "Open Product"})
                external_targets = {
                    relationship.attrib.get("Target")
                    for relationship in relationships.findall(f"{{{REL_NS}}}Relationship")
                    if relationship.attrib.get("TargetMode") == "External"
                }
                self.assertEqual(
                    external_targets,
                    {
                        "https://example.com/product-alpha",
                        "https://example.com/product-special",
                    },
                )

    def test_malformed_markdown_table_remains_plain_text(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            task_dir = Path(temporary) / "task"
            article = """# Malformed Table

This transition introduces an intentionally malformed table.

## Comparison

| Product | Quantity |
| --- | --- |
| Missing second cell |

The malformed block should remain editable text.

""" + FAQ_BLOCK
            task = SimpleNamespace(
                task_dir=str(task_dir),
                selected_title="Malformed Table",
                topic="Topic",
                article=article,
                final_article="",
                linked_article="",
                humanized_article="",
                initial_article="",
                hero_image="",
                products=[],
                images=[],
            )

            output = export_task_docx(test_config(), task)
            exported = Document(output)
            self.assertEqual(exported.tables, [])
            display = [paragraph_text(paragraph) for paragraph in exported.paragraphs]
            self.assertIn("| Product | Quantity |", display)
            self.assertIn("| Missing second cell |", display)

    def test_heading_links_and_bold_wrapped_links_render_as_real_hyperlinks(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            task_dir = Path(temporary) / "task"
            article = """# Linked Heading Export

This transition explains what buyers can compare.

## Product Options

### [Product Alpha](https://example.com/product-alpha)

**[Product Beta](https://example.com/product-beta)**

The product descriptions follow these links.

""" + FAQ_BLOCK
            task = SimpleNamespace(
                task_dir=str(task_dir),
                selected_title="Linked Heading Export",
                topic="Topic",
                article=article,
                final_article="",
                linked_article="",
                humanized_article="",
                initial_article="",
                hero_image="",
                products=[],
                images=[],
            )

            output = export_task_docx(test_config(), task)
            exported = Document(output)
            display = [paragraph_text(paragraph) for paragraph in exported.paragraphs]
            self.assertIn("Product Alpha", display)
            self.assertIn("Product Beta", display)
            self.assertNotIn("**", "\n".join(display))
            self.assertNotIn("[Product Alpha](https://example.com/product-alpha)", display)

            with zipfile.ZipFile(output) as archive:
                document_xml = ElementTree.fromstring(archive.read("word/document.xml"))
                relationships = ElementTree.fromstring(
                    archive.read("word/_rels/document.xml.rels")
                )
                relationship_targets = {
                    relationship.attrib["Id"]: relationship.attrib.get("Target")
                    for relationship in relationships.findall(f"{{{REL_NS}}}Relationship")
                    if relationship.attrib.get("TargetMode") == "External"
                }

                hyperlinks = document_xml.findall(f".//{{{WORD_NS}}}hyperlink")
                self.assertEqual(len(hyperlinks), 2)
                linked_text = {
                    "".join(
                        node.text or ""
                        for node in hyperlink.findall(f".//{{{WORD_NS}}}t")
                    ): relationship_targets[
                        hyperlink.attrib[
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
                        ]
                    ]
                    for hyperlink in hyperlinks
                }
                self.assertEqual(
                    linked_text,
                    {
                        "Product Alpha": "https://example.com/product-alpha",
                        "Product Beta": "https://example.com/product-beta",
                    },
                )

                alpha_hyperlink = next(
                    hyperlink
                    for hyperlink in hyperlinks
                    if "".join(
                        node.text or ""
                        for node in hyperlink.findall(f".//{{{WORD_NS}}}t")
                    ) == "Product Alpha"
                )
                alpha_run = alpha_hyperlink.find(f"{{{WORD_NS}}}r")
                self.assertIsNotNone(alpha_run.find(f"{{{WORD_NS}}}rPr/{{{WORD_NS}}}b"))
                self.assertEqual(
                    alpha_run.find(f"{{{WORD_NS}}}rPr/{{{WORD_NS}}}sz").attrib[
                        f"{{{WORD_NS}}}val"
                    ],
                    "26",
                )
                self.assertEqual(
                    alpha_run.find(f"{{{WORD_NS}}}rPr/{{{WORD_NS}}}color").attrib[
                        f"{{{WORD_NS}}}val"
                    ],
                    "0563C1",
                )
                self.assertEqual(
                    alpha_run.find(f"{{{WORD_NS}}}rPr/{{{WORD_NS}}}u").attrib[
                        f"{{{WORD_NS}}}val"
                    ],
                    "single",
                )

    def test_images_are_inline_marked_and_links_are_real_and_bold(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            task_dir = Path(temporary) / "task"
            hero = task_dir / "incoming" / "hero.png"
            product_image = task_dir / "incoming" / "product.png"
            make_png(hero, (30, 60, 90))
            make_png(product_image, (90, 120, 150))
            article = """# Inline Image Export

Read the [Company Website](https://company.example/) before comparing options.

## First Section

This section introduces the decision.

## Product Section

### Product Alpha

Choose [Product Alpha](https://example.com/product-alpha) for this application
when the confirmed dimensions match the assembly.

This qualification remains in the same product subsection.

## FAQ

**Q: What should buyers check first?**

A: Buyers should check the application requirements first.

**Q: When should buyers request a sample?**

A: Buyers should request one before approval when fit matters.

**Q: Why should buyers compare suppliers?**

A: Buyers should compare capability, quality control, delivery, and support.
"""
            product = SimpleNamespace(
                name="Product Alpha",
                url="https://example.com/product-alpha",
                image_path=str(product_image),
                description="",
            )
            task = SimpleNamespace(
                task_dir=str(task_dir),
                selected_title="Inline Image Export",
                topic="Topic",
                article=article,
                final_article="",
                linked_article="",
                humanized_article="",
                initial_article="",
                hero_image=str(hero),
                products=[product],
                images=[],
            )
            task.images = prepare_task_images(task, require_hero=True)

            output = export_task_docx(test_config(), task)
            exported = Document(output)
            paragraphs = exported.paragraphs
            display = ["[IMAGE]" if paragraph._p.xpath(".//w:drawing") else paragraph_text(paragraph) for paragraph in paragraphs]

            h1_index = display.index("Inline Image Export")
            transition_index = next(index for index, value in enumerate(display) if value.startswith("Read the Company"))
            hero_index = display.index("[IMAGE]")
            hero_marker_index = display.index("img.Inline Image Export.webp")
            first_h2_index = display.index("First Section")
            self.assertLess(h1_index, transition_index)
            self.assertLess(transition_index, hero_index)
            self.assertEqual(hero_marker_index, hero_index + 1)
            self.assertEqual(first_h2_index, hero_marker_index + 1)

            product_paragraph_index = next(
                index for index, value in enumerate(display) if value.startswith("Choose Product Alpha")
            )
            self.assertEqual(
                display[product_paragraph_index],
                "Choose Product Alpha for this application when the confirmed dimensions match the assembly.",
            )
            qualification_index = display.index(
                "This qualification remains in the same product subsection."
            )
            self.assertEqual(display[qualification_index + 1], "[IMAGE]")
            self.assertEqual(display[qualification_index + 2], "img.Product Alpha.webp")
            audit = (task_dir / "07_final_with_images.md").read_text(encoding="utf-8")
            self.assertLess(
                audit.index("This qualification remains in the same product subsection."),
                audit.index("img.Product Alpha.webp"),
            )
            self.assertLess(audit.index("img.Product Alpha.webp"), audit.index("## FAQ"))
            self.assertNotIn("Recommended Products", display)
            faq_index = display.index("FAQ")
            self.assertGreater(faq_index, product_paragraph_index)
            self.assertEqual(display[faq_index + 1], "Q: What should buyers check first?")
            self.assertNotIn("**", "\n".join(display))
            question_paragraph = paragraphs[faq_index + 1]
            self.assertTrue(question_paragraph.runs)
            self.assertTrue(all(run.bold for run in question_paragraph.runs if run.text))

            image_indices = [index for index, value in enumerate(display) if value == "[IMAGE]"]
            self.assertEqual(len(image_indices), 2)
            self.assertTrue(all(display[index + 1].startswith("img.") for index in image_indices))

            with zipfile.ZipFile(output) as archive:
                media = [name for name in archive.namelist() if name.startswith("word/media/")]
                self.assertEqual(len(media), 2)
                self.assertTrue(all(name.casefold().endswith(".webp") for name in media))
                content_types = archive.read("[Content_Types].xml").decode("utf-8")
                self.assertIn("image/webp", content_types)

                document_xml = ElementTree.fromstring(archive.read("word/document.xml"))
                for run in document_xml.findall(f".//{{{WORD_NS}}}r"):
                    fonts = run.find(f"{{{WORD_NS}}}rPr/{{{WORD_NS}}}rFonts")
                    self.assertIsNotNone(fonts)
                    self.assertEqual(
                        fonts.attrib[f"{{{WORD_NS}}}ascii"],
                        "Times New Roman",
                    )
                hyperlinks = document_xml.findall(f".//{{{WORD_NS}}}hyperlink")
                self.assertEqual(len(hyperlinks), 2)
                for hyperlink in hyperlinks:
                    self.assertIsNotNone(hyperlink.find(f".//{{{WORD_NS}}}b"))

                relationships = ElementTree.fromstring(
                    archive.read("word/_rels/document.xml.rels")
                )
                external_targets = {
                    relationship.attrib.get("Target")
                    for relationship in relationships.findall(f"{{{REL_NS}}}Relationship")
                    if relationship.attrib.get("TargetMode") == "External"
                }
                self.assertEqual(
                    external_targets,
                    {"https://company.example/", "https://example.com/product-alpha"},
                )

                heading_colors = []
                for paragraph in document_xml.findall(f".//{{{WORD_NS}}}p"):
                    style = paragraph.find(f"{{{WORD_NS}}}pPr/{{{WORD_NS}}}pStyle")
                    if style is None or style.attrib.get(f"{{{WORD_NS}}}val") not in {
                        "Heading1",
                        "Heading2",
                        "Heading3",
                    }:
                        continue
                    heading_colors.extend(
                        color.attrib.get(f"{{{WORD_NS}}}val")
                        for color in paragraph.findall(
                            f".//{{{WORD_NS}}}rPr/{{{WORD_NS}}}color"
                        )
                    )
                self.assertTrue(heading_colors)
                self.assertEqual(set(heading_colors), {"000000"})

    def test_hero_export_rejects_missing_transition(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            task_dir = Path(temporary) / "task"
            hero = task_dir / "hero.png"
            make_png(hero, (20, 20, 20))
            task = SimpleNamespace(
                task_dir=str(task_dir),
                selected_title="No Transition",
                topic="Topic",
                article=(
                    "# No Transition\n\n## First Section\n\nBody.\n\n" + FAQ_BLOCK
                ),
                hero_image=str(hero),
                products=[],
                images=[],
            )
            task.images = prepare_task_images(task, require_hero=True)

            with self.assertRaisesRegex(ArticleImageError, "过渡段"):
                export_task_docx(test_config(), task)

    def test_unresolved_product_image_requires_explicit_anchor(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            task_dir = Path(temporary) / "task"
            hero = task_dir / "hero.png"
            product_image = task_dir / "product.png"
            make_png(hero, (20, 40, 60))
            make_png(product_image, (70, 90, 110))
            task = SimpleNamespace(
                task_dir=str(task_dir),
                selected_title="Anchor Required",
                topic="Topic",
                article=(
                    "# Anchor Required\n\nA transition paragraph.\n\n"
                    "## First Section\n\nBody only.\n\n" + FAQ_BLOCK
                ),
                hero_image=str(hero),
                products=[
                    SimpleNamespace(
                        name="Missing Product",
                        url="https://example.com/missing",
                        image_path=str(product_image),
                    )
                ],
                images=[],
            )
            task.images = prepare_task_images(task, require_hero=True)

            with self.assertRaises(ImageAnchorRequiredError) as raised:
                export_task_docx(test_config(), task)
            self.assertEqual(
                [candidate["heading"] for candidate in raised.exception.unresolved[0]["anchor_candidates"]],
                ["First Section"],
            )


if __name__ == "__main__":
    unittest.main()
