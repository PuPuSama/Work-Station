from __future__ import annotations

import mimetypes
import re
import unicodedata
from collections.abc import Sequence
from datetime import date, datetime
from io import BytesIO
from pathlib import PurePath
from typing import Protocol, runtime_checkable
from zipfile import BadZipFile, ZipFile

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from .contracts import DocumentInput, ParsedAsset, ParsedBlock, ParsedDocument


class DocumentParserError(ValueError):
    """Base error for deterministic, user-safe document parsing failures."""


class UnsupportedDocumentError(DocumentParserError):
    """Raised when no registered parser accepts the supplied document."""


class DocumentParseError(DocumentParserError):
    """Raised when an accepted document cannot produce a normalized result."""


def _normalized_text(value: object) -> str:
    text = unicodedata.normalize("NFKC", str(value or ""))
    return re.sub(r"[ \t\f\v]+", " ", text).strip()


def _document_title(document: DocumentInput, candidate: str | None) -> str:
    normalized = _normalized_text(candidate)
    return normalized or PurePath(document.filename).stem


@runtime_checkable
class DocumentParser(Protocol):
    """Replaceable parser boundary that must return the normalized M2 contract."""

    @property
    def name(self) -> str: ...

    @property
    def version(self) -> str: ...

    def supports(self, document: DocumentInput) -> bool: ...

    def parse(self, document: DocumentInput) -> ParsedDocument: ...


class DocumentParserRouter:
    """Select exactly one parser from file extension and declared media type."""

    def __init__(self, parsers: Sequence[DocumentParser] | None = None) -> None:
        registered = tuple(parsers or default_document_parsers())
        if not registered:
            raise ValueError("at least one document parser is required")
        parser_names = tuple(parser.name for parser in registered)
        if len(set(parser_names)) != len(parser_names):
            raise ValueError("document parser names must be unique")
        self._parsers = registered

    @property
    def parsers(self) -> tuple[DocumentParser, ...]:
        return self._parsers

    def select(self, document: DocumentInput) -> DocumentParser:
        matches = tuple(parser for parser in self._parsers if parser.supports(document))
        if not matches:
            raise UnsupportedDocumentError(
                f"no document parser supports {document.suffix or 'this file type'}"
            )
        if len(matches) > 1:
            names = ", ".join(sorted(parser.name for parser in matches))
            raise DocumentParserError(
                f"document metadata matches multiple parsers: {names}"
            )
        return matches[0]

    def parse(self, document: DocumentInput) -> ParsedDocument:
        parser = self.select(document)
        parsed = parser.parse(document)
        if parsed.filename != document.filename:
            raise DocumentParseError("parser changed the source filename")
        if parsed.content_hash != document.content_hash:
            raise DocumentParseError("parser returned a mismatched source content hash")
        if parsed.parser_name != parser.name or parsed.parser_version != parser.version:
            raise DocumentParseError("parser returned an inconsistent identity")
        return parsed

    def close(self) -> None:
        for parser in self._parsers:
            close = getattr(parser, "close", None)
            if callable(close):
                close()


class DocxDocumentParser:
    name = "docx-lightweight"
    version = "1.0"
    _suffixes = frozenset({".docx"})
    _content_types = frozenset(
        {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
    )

    def supports(self, document: DocumentInput) -> bool:
        return (
            document.suffix in self._suffixes
            or document.content_type in self._content_types
        )

    def parse(self, document: DocumentInput) -> ParsedDocument:
        try:
            parsed = Document(BytesIO(document.content))
            blocks: list[ParsedBlock] = []
            heading_path: list[str] = []

            for paragraph_index, paragraph in enumerate(parsed.paragraphs):
                text = _normalized_text(paragraph.text)
                if not text:
                    continue
                style_name = _normalized_text(
                    getattr(getattr(paragraph, "style", None), "name", "")
                )
                heading_match = re.fullmatch(r"Heading ([1-9])", style_name)
                kind = "heading" if heading_match else "paragraph"
                if heading_match:
                    level = int(heading_match.group(1))
                    heading_path = heading_path[: level - 1]
                    heading_path.append(text)
                blocks.append(
                    ParsedBlock(
                        kind=kind,
                        ordinal=len(blocks),
                        text=text,
                        heading_path=tuple(heading_path),
                        locator={"paragraph_index": paragraph_index},
                        metadata={"style": style_name} if style_name else {},
                    )
                )

            for table_index, table in enumerate(parsed.tables):
                for row_index, row in enumerate(table.rows):
                    values = tuple(
                        _normalized_text(cell.text) for cell in row.cells
                    )
                    if not any(values):
                        continue
                    blocks.append(
                        ParsedBlock(
                            kind="table_row",
                            ordinal=len(blocks),
                            text="\t".join(values),
                            heading_path=tuple(heading_path),
                            locator={
                                "table_index": table_index,
                                "row_index": row_index,
                            },
                            metadata={"cell_count": len(values)},
                        )
                    )

            assets = self._embedded_assets(document)
            title = _document_title(document, parsed.core_properties.title)
            return ParsedDocument(
                filename=document.filename,
                content_type=next(iter(self._content_types)),
                content_hash=document.content_hash,
                parser_name=self.name,
                parser_version=self.version,
                blocks=tuple(blocks),
                assets=assets,
                title=title,
                metadata={
                    "paragraph_count": len(parsed.paragraphs),
                    "table_count": len(parsed.tables),
                },
            )
        except DocumentParserError:
            raise
        except (BadZipFile, KeyError, ValueError) as exc:
            raise DocumentParseError("DOCX document could not be parsed") from exc
        except Exception as exc:
            raise DocumentParseError("DOCX document could not be parsed") from exc

    @staticmethod
    def _embedded_assets(document: DocumentInput) -> tuple[ParsedAsset, ...]:
        assets: list[ParsedAsset] = []
        try:
            with ZipFile(BytesIO(document.content)) as archive:
                media_names = sorted(
                    name
                    for name in archive.namelist()
                    if name.startswith("word/media/") and not name.endswith("/")
                )
                for name in media_names:
                    content = archive.read(name)
                    if not content:
                        continue
                    filename = PurePath(name).name
                    content_type = (
                        mimetypes.guess_type(filename)[0]
                        or "application/octet-stream"
                    )
                    assets.append(
                        ParsedAsset(
                            filename=filename,
                            content=content,
                            content_type=content_type,
                            ordinal=len(assets),
                            locator={"package_path": name},
                        )
                    )
        except (BadZipFile, KeyError) as exc:
            raise DocumentParseError("DOCX embedded assets could not be read") from exc
        return tuple(assets)


class PdfDocumentParser:
    name = "pypdf-lightweight"
    version = "1.0"
    _suffixes = frozenset({".pdf"})
    _content_types = frozenset({"application/pdf"})

    def supports(self, document: DocumentInput) -> bool:
        return (
            document.suffix in self._suffixes
            or document.content_type in self._content_types
        )

    def parse(self, document: DocumentInput) -> ParsedDocument:
        try:
            reader = PdfReader(BytesIO(document.content))
            blocks: list[ParsedBlock] = []
            for page_index, page in enumerate(reader.pages):
                text = _normalized_text(page.extract_text() or "")
                if not text:
                    continue
                blocks.append(
                    ParsedBlock(
                        kind="page_text",
                        ordinal=len(blocks),
                        text=text,
                        locator={"page_number": page_index + 1},
                    )
                )
            title_value = None
            if reader.metadata is not None:
                title_value = getattr(reader.metadata, "title", None)
            return ParsedDocument(
                filename=document.filename,
                content_type="application/pdf",
                content_hash=document.content_hash,
                parser_name=self.name,
                parser_version=self.version,
                blocks=tuple(blocks),
                title=_document_title(document, title_value),
                metadata={"page_count": len(reader.pages)},
            )
        except Exception as exc:
            raise DocumentParseError(
                "PDF document contains no extractable text or could not be parsed"
            ) from exc


class ExcelDocumentParser:
    name = "openpyxl-lightweight"
    version = "1.0"
    _suffixes = frozenset({".xlsx", ".xlsm"})
    _content_types = frozenset(
        {
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel.sheet.macroenabled.12",
        }
    )

    def supports(self, document: DocumentInput) -> bool:
        return (
            document.suffix in self._suffixes
            or document.content_type in self._content_types
        )

    def parse(self, document: DocumentInput) -> ParsedDocument:
        workbook = None
        try:
            workbook = load_workbook(
                BytesIO(document.content),
                read_only=True,
                data_only=True,
            )
            blocks: list[ParsedBlock] = []
            for worksheet in workbook.worksheets:
                for row_number, row in enumerate(worksheet.iter_rows(values_only=True), 1):
                    values = tuple(self._cell_text(value) for value in row)
                    while values and not values[-1]:
                        values = values[:-1]
                    if not values or not any(values):
                        continue
                    blocks.append(
                        ParsedBlock(
                            kind="table_row",
                            ordinal=len(blocks),
                            text="\t".join(values),
                            heading_path=(worksheet.title,),
                            locator={
                                "sheet": worksheet.title,
                                "row_number": row_number,
                            },
                            metadata={"cell_count": len(values)},
                        )
                    )
            return ParsedDocument(
                filename=document.filename,
                content_type=(
                    "application/vnd.openxmlformats-officedocument."
                    "spreadsheetml.sheet"
                ),
                content_hash=document.content_hash,
                parser_name=self.name,
                parser_version=self.version,
                blocks=tuple(blocks),
                title=PurePath(document.filename).stem,
                metadata={"sheet_names": tuple(workbook.sheetnames)},
            )
        except Exception as exc:
            raise DocumentParseError(
                "Excel document contains no readable rows or could not be parsed"
            ) from exc
        finally:
            if workbook is not None:
                workbook.close()

    @staticmethod
    def _cell_text(value: object) -> str:
        if value is None:
            return ""
        if isinstance(value, (datetime, date)):
            return value.isoformat()
        if isinstance(value, bool):
            return "true" if value else "false"
        return _normalized_text(value)


def default_document_parsers() -> tuple[DocumentParser, ...]:
    return (
        DocxDocumentParser(),
        PdfDocumentParser(),
        ExcelDocumentParser(),
    )
